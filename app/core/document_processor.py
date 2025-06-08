"""
Document Processor Module

This module handles document uploads, text extraction, and processing for various file types.
It supports PDF, text, markdown, and other document formats.
"""

import os
import json
import uuid
import datetime
import logging
from typing import Dict, List, Optional, Union, BinaryIO

# Configure logging
logger = logging.getLogger(__name__)

# Try to import PDF processing libraries
try:
    import fitz  # PyMuPDF
    HAVE_PYMUPDF = True
except ImportError:
    HAVE_PYMUPDF = False
    logger.warning("PyMuPDF not available. PDF processing will be limited.")

# Try to import markdown
try:
    import markdown
    HAVE_MARKDOWN = True
except ImportError:
    HAVE_MARKDOWN = False
    logger.warning("Markdown library not available. Markdown processing will be limited.")

# Try to import docx
try:
    import docx
    HAVE_DOCX = True
    logger.info("python-docx library is available for DOCX processing")
except ImportError:
    HAVE_DOCX = False
    logger.warning("python-docx not available. DOCX processing will be limited.")

# Document storage path
DOCUMENTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "documents")

# Ensure documents directory exists
os.makedirs(DOCUMENTS_DIR, exist_ok=True)
logger.info(f"Document storage directory: {DOCUMENTS_DIR}")

def get_document_path(doc_id: str) -> str:
    """Get the file path for a document."""
    return os.path.join(DOCUMENTS_DIR, doc_id)

def process_uploaded_document(file_obj: BinaryIO, filename: str, character_id: str) -> Dict:
    """
    Process an uploaded document and extract its text.
    
    Args:
        file_obj: File-like object containing the document data
        filename: Original filename
        character_id: ID of the character to associate with the document
        
    Returns:
        Dictionary with document details
    """
    # Generate a unique ID for the document
    doc_id = str(uuid.uuid4())
    
    # Determine file extension
    _, ext = os.path.splitext(filename)
    ext = ext.lower()
    
    logger.info(f"Processing document: {filename}, extension: {ext}, id: {doc_id}")
    
    # Save the original file
    file_path = get_document_path(f"{doc_id}{ext}")
    with open(file_path, 'wb') as f:
        f.write(file_obj.read())
    
    logger.info(f"Saved document to: {file_path}")
    
    # Extract text based on file type
    extracted_text = ""
    doc_type = "unknown"
    
    if ext == '.pdf' and HAVE_PYMUPDF:
        logger.info("Processing PDF document")
        extracted_text = extract_text_from_pdf(file_path)
        doc_type = "pdf"
    elif ext == '.txt':
        logger.info("Processing TXT document")
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            extracted_text = f.read()
        doc_type = "text"
    elif ext in ['.md', '.markdown'] and HAVE_MARKDOWN:
        logger.info("Processing Markdown document")
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            md_text = f.read()
            extracted_text = md_text  # Store original markdown
        doc_type = "markdown"
    elif ext in ['.docx', '.doc'] and HAVE_DOCX:
        logger.info("Processing DOCX document")
        extracted_text = extract_text_from_docx(file_path)
        doc_type = "docx"
    else:
        logger.warning(f"Unsupported document format: {ext}")
        # For unsupported formats, just store metadata
        extracted_text = f"Unsupported document format: {ext}"
        doc_type = "unsupported"
    
    logger.info(f"Extracted {len(extracted_text)} characters from document")
    
    # Create document metadata
    document = {
        "id": doc_id,
        "filename": filename,
        "character_id": character_id,
        "file_path": file_path,
        "doc_type": doc_type,
        "text_content": extracted_text,
        "upload_date": datetime.datetime.now().isoformat(),
        "file_size": os.path.getsize(file_path)
    }
    
    # Save metadata
    metadata_path = get_document_path(f"{doc_id}.json")
    with open(metadata_path, 'w', encoding='utf-8') as f:
        json.dump(document, f, indent=2)
    
    logger.info(f"Document metadata saved to: {metadata_path}")
    return document

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text
    """
    if not HAVE_PYMUPDF:
        return "PDF text extraction requires PyMuPDF library."
    
    try:
        text = ""
        with fitz.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return f"Error extracting text: {str(e)}"

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Extracted text
    """
    if not HAVE_DOCX:
        logger.error("Attempted to process DOCX but python-docx is not available")
        return "DOCX text extraction requires python-docx library."
    
    try:
        logger.info(f"Opening DOCX file: {docx_path}")
        doc = docx.Document(docx_path)
        full_text = []
        
        # Extract text from paragraphs
        logger.info(f"Extracting text from {len(doc.paragraphs)} paragraphs")
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        logger.info(f"Extracting text from {len(doc.tables)} tables")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        # Join all text with newlines
        result = '\n'.join(full_text)
        logger.info(f"Successfully extracted {len(result)} characters from DOCX")
        return result
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return f"Error extracting text: {str(e)}"

def get_document_by_id(doc_id: str) -> Optional[Dict]:
    """
    Get document details by ID.
    
    Args:
        doc_id: Document ID
        
    Returns:
        Document details dictionary or None if not found
    """
    metadata_path = get_document_path(f"{doc_id}.json")
    if not os.path.exists(metadata_path):
        logger.warning(f"Document not found: {doc_id}")
        return None
    
    try:
        with open(metadata_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"Error reading document metadata: {e}")
        return None

def get_character_documents(character_id: str, limit: int = 20, offset: int = 0) -> List[Dict]:
    """
    Get documents associated with a character.
    
    Args:
        character_id: Character ID
        limit: Maximum number of documents to return
        offset: Number of documents to skip
        
    Returns:
        List of document details
    """
    documents = []
    try:
        for filename in os.listdir(DOCUMENTS_DIR):
            if not filename.endswith('.json'):
                continue
            
            try:
                with open(os.path.join(DOCUMENTS_DIR, filename), 'r', encoding='utf-8') as f:
                    doc = json.load(f)
                    if doc.get('character_id') == character_id:
                        documents.append(doc)
            except Exception as e:
                logger.error(f"Error reading document {filename}: {e}")
                continue
        
        # Sort by upload date (newest first)
        documents.sort(key=lambda x: x.get('upload_date', ''), reverse=True)
        
        # Apply pagination
        return documents[offset:offset + limit]
    except Exception as e:
        logger.error(f"Error listing character documents: {e}")
        return []

def delete_document(doc_id: str) -> bool:
    """
    Delete a document and its metadata.
    
    Args:
        doc_id: Document ID
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Get document details first
        doc = get_document_by_id(doc_id)
        if not doc:
            logger.warning(f"Cannot delete non-existent document: {doc_id}")
            return False
        
        # Delete the original file
        if os.path.exists(doc['file_path']):
            os.remove(doc['file_path'])
            logger.info(f"Deleted original file: {doc['file_path']}")
        
        # Delete metadata
        metadata_path = get_document_path(f"{doc_id}.json")
        if os.path.exists(metadata_path):
            os.remove(metadata_path)
            logger.info(f"Deleted metadata file: {metadata_path}")
        
        return True
    except Exception as e:
        logger.error(f"Error deleting document {doc_id}: {e}")
        return False

def create_document_summary(doc_id: str, character_id: str) -> str:
    """
    Create a summary of a document for a character.
    
    Args:
        doc_id: Document ID
        character_id: Character ID
        
    Returns:
        Summary text
    """
    try:
        doc = get_document_by_id(doc_id)
        if not doc:
            logger.warning(f"Cannot summarize non-existent document: {doc_id}")
            return "Document not found."
        
        if doc.get('character_id') != character_id:
            logger.warning(f"Document {doc_id} does not belong to character {character_id}")
            return "Document not found."
        
        # For now, just return the first 500 characters
        text = doc.get('text_content', '')
        return text[:500] + ('...' if len(text) > 500 else '')
    except Exception as e:
        logger.error(f"Error creating document summary: {e}")
        return "Error creating summary."

def extract_document_to_memory(doc_id: str, character_id: str) -> bool:
    """
    Extract document content to character's memory.
    
    Args:
        doc_id: Document ID
        character_id: Character ID
        
    Returns:
        True if successful, False otherwise
    """
    try:
        doc = get_document_by_id(doc_id)
        if not doc:
            logger.warning(f"Cannot extract non-existent document: {doc_id}")
            return False
        
        if doc.get('character_id') != character_id:
            logger.warning(f"Document {doc_id} does not belong to character {character_id}")
            return False
        
        # For now, just return success
        # TODO: Implement actual memory extraction
        return True
    except Exception as e:
        logger.error(f"Error extracting document to memory: {e}")
        return False 