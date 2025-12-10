#!/usr/bin/env python3
"""
Create sample test documents for demonstrating the RAG system.
"""

import os
import pandas as pd
from docx import Document as DocxDocument
from pathlib import Path

def create_test_documents():
    """Create a variety of test documents for RAG demonstration."""
    
    # Create output directory
    output_dir = Path(__file__).parent / "documents"
    output_dir.mkdir(exist_ok=True)
    
    print(f"Creating test documents in: {output_dir}")
    
    # 1. Business Plan (Text file)
    business_plan = """
# MegaCorp AI Solutions - Business Plan 2024

## Executive Summary
MegaCorp AI Solutions is a cutting-edge technology company specializing in privacy-first artificial intelligence solutions. Our mission is to democratize AI while maintaining complete user privacy and data sovereignty.

## Company Overview
- **Founded**: 2024
- **Headquarters**: San Francisco, CA
- **Team Size**: 25 employees
- **Focus**: Privacy-preserving AI, Local LLM deployment, Enterprise AI solutions

## Market Analysis
The global AI market is projected to reach $1.8 trillion by 2030, growing at a CAGR of 38.1%. Key trends include:
- Increased demand for privacy-compliant AI solutions
- Enterprise adoption of local AI deployment
- Growing concern over data sovereignty
- Regulatory pressure for transparent AI systems

## Products & Services

### Core Products
1. **PrivateChat Pro** - On-premises conversational AI platform
2. **LocalLLM Enterprise** - Self-hosted large language model solutions  
3. **AI Audit Suite** - Compliance and transparency tools
4. **SecureRAG** - Privacy-first retrieval-augmented generation

### Target Markets
- **Healthcare**: HIPAA-compliant AI assistants
- **Finance**: SOX-compliant document analysis
- **Government**: Classified information processing
- **Legal**: Attorney-client privilege preserving AI

## Financial Projections

### Revenue Targets
- **Year 1 (2024)**: $2.5 million
- **Year 2 (2025)**: $8.5 million  
- **Year 3 (2026)**: $25 million
- **Year 4 (2027)**: $60 million
- **Year 5 (2028)**: $120 million

### Key Metrics
- **Customer Acquisition Cost (CAC)**: $15,000
- **Customer Lifetime Value (CLV)**: $125,000
- **Monthly Recurring Revenue (MRR) Growth**: 25%
- **Gross Margin**: 85%
- **Net Profit Margin Target**: 35% by Year 3

## Competitive Advantage
1. **Privacy-First Architecture**: All processing happens on-premises
2. **Regulatory Compliance**: Built-in GDPR, HIPAA, SOX compliance
3. **Transparent AI**: Full explainability and audit trails
4. **Cost Efficiency**: 70% lower operating costs vs cloud solutions
5. **Data Sovereignty**: Complete customer control over data

## Team & Leadership
- **CEO**: Sarah Chen, Former VP of AI at Google
- **CTO**: Dr. Marcus Rodriguez, PhD in Distributed Systems  
- **CPO**: Jennifer Liu, 15 years in enterprise software
- **Head of Sales**: Robert Kim, Former Director at Salesforce

## Funding Requirements
**Seeking Series A**: $15 million
- **Product Development**: $8M (53%)
- **Sales & Marketing**: $4M (27%)
- **Operations**: $2M (13%)
- **Legal & Compliance**: $1M (7%)

## Risk Factors
1. **Technical Risk**: Complexity of on-premises AI deployment
2. **Market Risk**: Slower enterprise adoption cycles
3. **Competitive Risk**: Big Tech entering privacy-focused AI
4. **Regulatory Risk**: Changing AI compliance requirements

## Milestones
- **Q1 2024**: Complete Series A funding
- **Q2 2024**: Launch PrivateChat Pro beta
- **Q3 2024**: First enterprise customer (Fortune 500)
- **Q4 2024**: $1M ARR milestone
- **Q1 2025**: International expansion (EU market)

## Contact Information
- **Email**: investors@megacorp-ai.com
- **Phone**: +1 (555) 123-4567
- **Website**: www.megacorp-ai.com
- **Address**: 123 Innovation Drive, San Francisco, CA 94107
    """
    
    with open(output_dir / "business_plan.txt", "w") as f:
        f.write(business_plan)
    
    # 2. Employee Handbook (Word document)
    handbook = DocxDocument()
    handbook.add_heading('MegaCorp Employee Handbook', 0)
    
    handbook.add_heading('Welcome Message', level=1)
    handbook.add_paragraph(
        'Welcome to MegaCorp AI Solutions! This handbook contains important information '
        'about company policies, benefits, and procedures. Please read it carefully and '
        'keep it handy for reference.'
    )
    
    handbook.add_heading('Company Policies', level=1)
    
    handbook.add_heading('Remote Work Policy', level=2)
    handbook.add_paragraph('MegaCorp supports flexible work arrangements:')
    p = handbook.add_paragraph()
    p.add_run('• ').bold = True
    p.add_run('Hybrid Schedule: 3 days in office, 2 days remote')
    p = handbook.add_paragraph()
    p.add_run('• ').bold = True
    p.add_run('Mandatory Office Days: Tuesday and Thursday')
    p = handbook.add_paragraph()
    p.add_run('• ').bold = True
    p.add_run('Core Hours: 10 AM - 3 PM (all team members available)')
    p = handbook.add_paragraph()
    p.add_run('• ').bold = True
    p.add_run('Equipment: Company provides laptop, monitor, and $500 home office stipend')
    
    handbook.add_heading('Vacation Policy', level=2)
    handbook.add_paragraph('MegaCorp values work-life balance:')
    handbook.add_paragraph('• Unlimited PTO policy (minimum 3 weeks recommended)')
    handbook.add_paragraph('• All major holidays observed')
    handbook.add_paragraph('• Company shutdown: December 25 - January 2')
    handbook.add_paragraph('• Sabbatical: 4 weeks after 5 years of service')
    
    handbook.add_heading('Professional Development', level=2)
    handbook.add_paragraph('Investment in employee growth:')
    handbook.add_paragraph('• Annual learning budget: $2,500 per employee')
    handbook.add_paragraph('• Conference attendance: 1 major conference per year')
    handbook.add_paragraph('• Internal tech talks: Every Friday afternoon')
    handbook.add_paragraph('• Mentorship program: Mandatory for all levels')
    
    handbook.add_heading('Health & Benefits', level=2)
    handbook.add_paragraph('Comprehensive benefits package:')
    handbook.add_paragraph('• Health Insurance: 100% premium covered (employee + family)')
    handbook.add_paragraph('• Dental & Vision: 100% premium covered')
    handbook.add_paragraph('• Mental Health: $150/month counseling allowance')
    handbook.add_paragraph('• Gym Membership: $100/month fitness reimbursement')
    handbook.add_paragraph('• Life Insurance: 2x annual salary')
    
    handbook.add_heading('Code of Conduct', level=2)
    handbook.add_paragraph('Our values guide our behavior:')
    handbook.add_paragraph('• Privacy First: Respect user data and privacy always')
    handbook.add_paragraph('• Transparency: Open communication and honest feedback')
    handbook.add_paragraph('• Inclusion: Diverse perspectives make us stronger')
    handbook.add_paragraph('• Excellence: Deliver high-quality work consistently')
    handbook.add_paragraph('• Innovation: Challenge the status quo respectfully')
    
    handbook.save(output_dir / "employee_handbook.docx")
    
    # 3. Financial Data (Excel/CSV)
    financial_data = {
        'Month': ['Jan 2024', 'Feb 2024', 'Mar 2024', 'Apr 2024', 'May 2024', 'Jun 2024',
                 'Jul 2024', 'Aug 2024', 'Sep 2024', 'Oct 2024', 'Nov 2024', 'Dec 2024'],
        'Revenue': [150000, 185000, 210000, 240000, 280000, 320000,
                   365000, 410000, 455000, 500000, 550000, 600000],
        'Expenses': [120000, 135000, 150000, 165000, 180000, 195000,
                    210000, 225000, 240000, 255000, 270000, 285000],
        'Net_Profit': [30000, 50000, 60000, 75000, 100000, 125000,
                      155000, 185000, 215000, 245000, 280000, 315000],
        'Customers': [25, 32, 41, 52, 64, 78, 94, 112, 131, 152, 175, 200],
        'MRR': [125000, 155000, 175000, 200000, 233000, 267000,
               304000, 342000, 379000, 417000, 458000, 500000]
    }
    
    df = pd.DataFrame(financial_data)
    df.to_csv(output_dir / "financial_data.csv", index=False)
    
    # Also create Excel version with multiple sheets
    with pd.ExcelWriter(output_dir / "financial_report.xlsx") as writer:
        df.to_excel(writer, sheet_name='Monthly_Summary', index=False)
        
        # Product breakdown
        product_data = {
            'Product': ['PrivateChat Pro', 'LocalLLM Enterprise', 'AI Audit Suite', 'SecureRAG'],
            'Q1_Revenue': [75000, 45000, 15000, 15000],
            'Q2_Revenue': [145000, 85000, 35000, 35000],
            'Q3_Revenue': [220000, 130000, 60000, 50000],
            'Q4_Revenue': [295000, 175000, 85000, 65000],
            'Total_Customers': [120, 45, 25, 15],
            'Avg_Deal_Size': [12500, 22000, 8500, 6500]
        }
        pd.DataFrame(product_data).to_excel(writer, sheet_name='Product_Breakdown', index=False)
    
    # 4. Technical Documentation (Markdown)
    tech_docs = """
# MegaCorp AI Technical Architecture

## System Overview
MegaCorp's AI platform is built on a microservices architecture designed for privacy, scalability, and compliance.

## Core Components

### 1. Privacy Engine
- **Zero-Trust Architecture**: No data leaves customer premises
- **Homomorphic Encryption**: Computation on encrypted data
- **Differential Privacy**: Statistical privacy guarantees
- **Secure Enclaves**: Hardware-based isolation

### 2. LLM Management
- **Model Serving**: Optimized inference engines
- **Model Registry**: Version control and deployment
- **Resource Management**: Dynamic GPU allocation
- **Performance Monitoring**: Real-time metrics

### 3. RAG System
- **Vector Database**: FAISS with custom indexing
- **Embedding Models**: sentence-transformers optimized
- **Chunk Management**: Smart document segmentation
- **Context Ranking**: Relevance-based retrieval

## Security Framework

### Data Protection
- **Encryption at Rest**: AES-256 encryption
- **Encryption in Transit**: TLS 1.3
- **Key Management**: Hardware Security Modules (HSM)
- **Access Control**: Role-based permissions

### Compliance Standards
- **GDPR**: Right to deletion, data portability
- **HIPAA**: PHI protection, audit logs
- **SOX**: Financial data controls
- **ISO 27001**: Information security management

## Performance Specifications

### Latency Targets
- **Chat Response**: < 500ms (p95)
- **Document Processing**: < 30s per MB
- **Search Queries**: < 100ms (p99)
- **Model Loading**: < 60s cold start

### Throughput Capacity
- **Concurrent Users**: 10,000+ per instance
- **Document Ingestion**: 1TB per hour
- **API Requests**: 100,000 requests/minute
- **Vector Search**: 1M queries/second

## Deployment Architecture

### On-Premises Deployment
```
┌─────────────────┐    ┌─────────────────┐    ┌─────────────────┐
│   Load Balancer │    │   API Gateway   │    │  Auth Service   │
└─────────────────┘    └─────────────────┘    └─────────────────┘
         │                       │                       │
         └───────────────────────┼───────────────────────┘
                                 │
         ┌───────────────────────┼───────────────────────┐
         │                       │                       │
┌─────────────────┐    ┌─────────────────┐    ┌─────────────────┐
│  Chat Service   │    │Document Service │    │  Vector DB      │
└─────────────────┘    └─────────────────┘    └─────────────────┘
         │                       │                       │
         └───────────────────────┼───────────────────────┘
                                 │
                    ┌─────────────────┐
                    │   LLM Engine    │
                    └─────────────────┘
```

### Cloud Hybrid Option
- **Control Plane**: Customer's cloud (metadata only)
- **Data Plane**: On-premises (sensitive data)
- **Model Updates**: Encrypted, signed packages
- **Monitoring**: Privacy-preserving telemetry

## API Specifications

### Authentication
- **OAuth 2.0**: Standard authentication flow
- **JWT Tokens**: Stateless authentication
- **API Keys**: Service-to-service auth
- **MFA Support**: Multi-factor authentication

### Rate Limiting
- **Per User**: 1,000 requests/hour
- **Per Organization**: 100,000 requests/hour
- **Burst Allowance**: 2x sustained rate for 5 minutes
- **Enterprise**: Custom limits available

## Monitoring & Observability

### Metrics Collection
- **System Metrics**: CPU, memory, disk, network
- **Application Metrics**: Response times, error rates
- **Business Metrics**: Usage patterns, feature adoption
- **Security Metrics**: Failed auth attempts, anomalies

### Alerting
- **Performance**: Response time > 1s
- **Availability**: Service down > 30s
- **Security**: Suspicious activity detected
- **Capacity**: Resource utilization > 80%

## Disaster Recovery

### Backup Strategy
- **Real-time Replication**: Cross-zone redundancy
- **Daily Snapshots**: Point-in-time recovery
- **Weekly Archives**: Long-term retention
- **Offsite Storage**: Geographically distributed

### Recovery Objectives
- **RTO (Recovery Time)**: < 4 hours
- **RPO (Recovery Point)**: < 15 minutes
- **Availability**: 99.9% uptime SLA
- **Data Integrity**: Zero data loss guarantee
    """
    
    with open(output_dir / "technical_documentation.md", "w") as f:
        f.write(tech_docs)
    
    # 5. Meeting Notes (Text file)
    meeting_notes = """
# Weekly Team Meeting Notes - December 15, 2024

**Attendees**: Sarah Chen (CEO), Marcus Rodriguez (CTO), Jennifer Liu (CPO), Robert Kim (Sales)
**Time**: 10:00 AM - 11:00 AM PST
**Location**: Conference Room A / Zoom Hybrid

## Action Items from Last Week
✅ Complete security audit for PrivateChat Pro (Marcus)
✅ Finalize pricing strategy for Q1 launch (Jennifer)  
✅ Set up demo environment for prospect meetings (Robert)
🔄 Hire senior ML engineer - 3 candidates in final round (Sarah)

## Key Updates

### Product Development (Marcus)
- **PrivateChat Pro**: Beta testing with 5 enterprise customers
  - Performance improvements: 40% faster response times
  - Bug fixes: Resolved 12 critical issues
  - New features: Multi-language support, advanced search
  - **Concern**: Memory usage higher than expected with large documents
  - **Action**: Optimize embedding storage by Jan 15

- **SecureRAG**: Alpha version ready for internal testing
  - Vector search performance: 95% queries under 100ms
  - Document processing: Supports 15 file formats
  - **Issue**: PDF extraction accuracy needs improvement
  - **Action**: Integrate better PDF parser by Dec 30

### Sales & Marketing (Robert)
- **Pipeline Update**: $2.3M in qualified opportunities
  - 15 active prospects in various stages
  - 3 POCs (Proof of Concepts) scheduled for January
  - 1 enterprise deal ($750K) in legal review
  
- **Target Sectors Performance**:
  - Healthcare: Strong interest, compliance questions
  - Finance: Price sensitivity, need ROI justification  
  - Legal: High engagement, long sales cycles
  - Government: Security focus, budget constraints

- **Marketing Initiatives**:
  - Conference speaking: RSA 2024, AI Security Summit
  - Content marketing: 2 whitepapers published
  - Partnership: Exploring integration with Microsoft Teams

### Product Strategy (Jennifer)
- **Customer Feedback Analysis**: 85% satisfaction rate
  - **Top Requests**: 
    1. Mobile app (60% of customers)
    2. Advanced analytics dashboard (45%)
    3. API rate limit increases (40%)
    4. Integration with popular tools (35%)
  
- **Competitive Analysis**:
  - OpenAI Enterprise: Gaining market share but privacy concerns
  - Anthropic: Strong technical capabilities, limited enterprise features
  - Google Cloud AI: Comprehensive but complex deployment
  - **Our Advantage**: Privacy-first approach resonating well

### Business Operations (Sarah)
- **Funding Status**: Series A progressing well
  - 3 VCs showing strong interest
  - Term sheet expected by January 31
  - Current runway: 8 months at current burn rate
  
- **Team Growth**:
  - Engineering: Need 3 senior developers
  - Sales: Adding 2 AEs (Account Executives)
  - Marketing: Hiring content marketing manager
  - **Challenge**: Competitive talent market in AI space

## Key Decisions Made

1. **Pricing Strategy**: Finalized tiered pricing for PrivateChat Pro
   - Starter: $50/user/month (up to 100 users)
   - Professional: $150/user/month (enterprise features)
   - Enterprise: Custom pricing (dedicated deployment)

2. **Go-to-Market Focus**: Prioritize healthcare and finance verticals
   - Higher willingness to pay for privacy
   - Clear compliance requirements
   - Shorter evaluation cycles than government

3. **Partnership Strategy**: Pursue system integrator partnerships
   - Target: Deloitte, PwC, EY consulting practices
   - Goal: 3 partnerships signed by Q2 2024

## Challenges & Risks

### Technical Challenges
- **Scaling**: Current architecture limits to 1,000 concurrent users
- **Model Performance**: Need better accuracy for specialized domains
- **Integration**: Complex enterprise IT environments

### Market Challenges  
- **Education**: Market needs education on privacy benefits
- **Competition**: Big Tech companies entering space
- **Regulation**: Unclear AI regulation landscape

### Operational Challenges
- **Talent Acquisition**: Difficult to find AI experts
- **Customer Support**: Need specialized technical support team
- **Compliance**: Multiple industry standards to maintain

## Next Week Priorities

### Engineering
- [ ] Complete memory optimization for large documents
- [ ] Start mobile app prototype development  
- [ ] Security penetration testing with external firm

### Sales
- [ ] Conduct 3 customer POCs
- [ ] Finalize enterprise deal in legal review
- [ ] Attend Healthcare AI Conference in Boston

### Product
- [ ] User research interviews with 10 beta customers
- [ ] Complete competitive analysis report
- [ ] Define Q1 feature roadmap

### Operations
- [ ] Interview final ML engineer candidates
- [ ] Prepare Series A data room
- [ ] Legal review of partnership agreements

## Metrics Dashboard

### Product Metrics
- Active Users: 2,847 (+15% WoW)
- API Calls: 1.2M (+25% WoW)  
- Uptime: 99.94% (target: 99.9%)
- Support Tickets: 23 (-30% WoW)

### Business Metrics
- MRR: $485K (+12% MoM)
- Customer Count: 187 (+8 new this week)
- Churn Rate: 2.1% (target: <3%)
- NPS Score: 72 (target: >70)

**Next Meeting**: December 22, 2024 at 10:00 AM PST
**Note**: No meeting December 29 (holiday week)
    """
    
    with open(output_dir / "meeting_notes.txt", "w") as f:
        f.write(meeting_notes)
    
    # 6. Customer Support FAQ (Text file)
    faq_content = """
# Customer Support FAQ - MegaCorp AI Solutions

## General Questions

### Q: What makes MegaCorp different from other AI companies?
A: MegaCorp is privacy-first. All AI processing happens on your premises or private cloud. Your data never leaves your control, ensuring complete privacy and compliance with regulations like GDPR and HIPAA.

### Q: What industries do you serve?
A: We specialize in highly regulated industries including healthcare, finance, legal, and government. Our solutions are designed to meet strict compliance requirements while delivering powerful AI capabilities.

### Q: Do you offer cloud-based solutions?
A: We offer hybrid solutions where the control plane can be in your preferred cloud, but all sensitive data processing happens on-premises. This gives you cloud convenience with on-premises privacy.

## Technical Questions

### Q: What file formats does SecureRAG support?
A: SecureRAG supports 15+ file formats including:
- Documents: PDF, DOCX, DOC, TXT, MD, RTF
- Spreadsheets: XLSX, XLS, CSV
- Presentations: PPTX, PPT
- Images: PNG, JPG (with OCR)
- Code: Multiple programming languages
- Structured data: JSON, XML, YAML

### Q: How fast is document processing?
A: Processing speed depends on file size and content:
- Text files: < 5 seconds per MB
- PDFs: < 30 seconds per MB
- Images with OCR: < 60 seconds per MB
- Large spreadsheets: < 45 seconds per MB

### Q: What's the maximum file size supported?
A: Standard limits:
- Individual files: 100 MB
- Batch upload: 1 GB
- Enterprise customers: Custom limits available
- Storage: Unlimited (customer infrastructure)

### Q: How accurate is the AI?
A: Accuracy varies by use case:
- Document Q&A: 92% accuracy
- Content summarization: 89% accuracy
- Data extraction: 94% accuracy
- Code analysis: 87% accuracy

### Q: What LLM models do you support?
A: We support various models:
- **Open source**: Llama 2/3, Mistral, Code Llama
- **Commercial**: GPT-4, Claude (via secure API)
- **Specialized**: Domain-specific fine-tuned models
- **Custom**: We can help you fine-tune models

## Security & Compliance

### Q: How do you ensure data privacy?
A: Multiple layers of protection:
- **Encryption**: AES-256 at rest, TLS 1.3 in transit
- **Access Control**: Role-based permissions, MFA
- **Audit Logging**: Complete activity tracking
- **Zero Trust**: No data leaves your environment
- **Compliance**: GDPR, HIPAA, SOX certified

### Q: What compliance certifications do you have?
A: Current certifications:
- SOC 2 Type II
- ISO 27001
- GDPR compliant
- HIPAA compliant
- SOX controls available
- FedRAMP (in progress)

### Q: How do you handle data breaches?
A: Since data never leaves your premises, traditional cloud breaches aren't a concern. However, we provide:
- Incident response procedures
- Security monitoring tools
- Breach notification systems
- Forensic analysis capabilities
- Recovery procedures

## Pricing & Licensing

### Q: How is pricing structured?
A: Flexible pricing models:
- **Per User**: $50-$150/user/month
- **Per Document**: $0.10-$0.50 per document processed
- **Flat Rate**: $10K-$50K/month unlimited usage
- **Enterprise**: Custom pricing for large deployments

### Q: Are there setup fees?
A: Setup fees vary:
- **Self-service**: No setup fee
- **Assisted setup**: $5,000-$15,000
- **Full deployment**: $25,000-$100,000
- **Custom integration**: Quote-based

### Q: What's included in support?
A: Support tiers:
- **Standard**: Email support, knowledge base
- **Professional**: Phone support, faster response
- **Enterprise**: Dedicated support engineer, 24/7 coverage
- **Premium**: On-site support, custom training

## Implementation

### Q: How long does implementation take?
A: Timeline depends on complexity:
- **Simple deployment**: 1-2 weeks
- **Standard enterprise**: 4-6 weeks  
- **Complex integration**: 8-12 weeks
- **Custom development**: 3-6 months

### Q: What infrastructure do I need?
A: Minimum requirements:
- **CPU**: 16 cores
- **RAM**: 64 GB
- **Storage**: 1 TB SSD
- **GPU**: Optional but recommended (RTX 4090 or better)
- **Network**: 1 Gbps
- **OS**: Linux (Ubuntu 20.04+) or Windows Server 2019+

### Q: Do you provide training?
A: Comprehensive training options:
- **Online training**: Self-paced modules
- **Live webinars**: Weekly product demos
- **On-site training**: Custom workshops
- **Certification program**: Technical certification
- **Documentation**: Complete user guides

## Integration

### Q: What systems do you integrate with?
A: Popular integrations:
- **CRM**: Salesforce, HubSpot, Microsoft Dynamics
- **Productivity**: Microsoft 365, Google Workspace
- **Collaboration**: Slack, Teams, Zoom
- **Storage**: SharePoint, Box, Dropbox
- **Security**: Okta, Active Directory, LDAP

### Q: Do you have APIs?
A: Comprehensive API coverage:
- **RESTful APIs**: Full functionality access
- **WebSocket**: Real-time streaming
- **GraphQL**: Flexible data queries
- **Webhooks**: Event-driven integrations
- **SDKs**: Python, JavaScript, Java, C#

### Q: Can I customize the interface?
A: Multiple customization options:
- **Branding**: Logo, colors, themes
- **Layout**: Custom dashboards
- **Features**: Enable/disable functionality
- **Workflows**: Custom business processes
- **White-label**: Complete rebranding available

## Troubleshooting

### Q: What if processing fails?
A: Troubleshooting steps:
1. Check file format and size
2. Verify system resources
3. Review error logs
4. Retry with different settings
5. Contact support if issues persist

### Q: Why are responses slow?
A: Common causes and solutions:
- **Large documents**: Process in smaller chunks
- **Limited resources**: Upgrade hardware
- **Network issues**: Check connectivity
- **High usage**: Scale infrastructure
- **Model selection**: Use faster models

### Q: How do I update the system?
A: Update process:
1. **Notification**: We'll alert you to updates
2. **Staging**: Test in staging environment
3. **Backup**: Automatic backup before update
4. **Update**: Typically takes 30-60 minutes
5. **Verification**: Confirm system functionality

## Contact Information

### Support Channels
- **Email**: support@megacorp-ai.com
- **Phone**: +1 (555) 123-4567
- **Chat**: Available 9 AM - 6 PM PST
- **Portal**: https://support.megacorp-ai.com

### Response Times
- **Critical**: 1 hour
- **High**: 4 hours
- **Medium**: 24 hours
- **Low**: 72 hours

### Escalation
For urgent issues:
- **Phone**: +1 (555) 123-4568 (24/7)
- **Emergency**: escalation@megacorp-ai.com
- **Account Manager**: Contact your dedicated AM

---
*Last updated: December 15, 2024*
*Version: 2.1*
    """
    
    with open(output_dir / "customer_support_faq.txt", "w") as f:
        f.write(faq_content)
    
    print("✅ Test documents created successfully!")
    print("\nCreated files:")
    for file in output_dir.glob("*"):
        print(f"  - {file.name} ({file.stat().st_size:,} bytes)")
    
    print(f"\n📁 Upload these files to test the RAG system:")
    print(f"   Navigate to http://localhost:8080/documents")
    print(f"   Upload files from: {output_dir}")
    
    print(f"\n🤖 Test queries to try after uploading:")
    print("   • 'What is our revenue target for 2025?'")
    print("   • 'What's the remote work policy?'") 
    print("   • 'How many vacation days do employees get?'")
    print("   • 'What are the system requirements?'")
    print("   • 'Who should I contact for support?'")
    print("   • 'What file formats are supported?'")

if __name__ == "__main__":
    create_test_documents()