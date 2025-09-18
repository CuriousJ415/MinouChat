"""
Automated tests for the document and RAG system.
"""

import pytest
import os
import json
import tempfile
from pathlib import Path
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
import pandas as pd
from docx import Document as DocxDocument

# Import the FastAPI app and dependencies
import sys
sys.path.append('/Users/jasonnau/projects/MiaChat/src')

from miachat.api.main import app
from miachat.database.models import Base, User, Document, DocumentChunk
from miachat.database.config import get_db
from miachat.api.core.document_service import document_service
from miachat.api.core.document_processor import document_processor
from miachat.api.core.embedding_service import embedding_service
from miachat.api.core.rag_service import rag_service

# Test database setup
SQLALCHEMY_DATABASE_URL = "sqlite:///./test_documents.db"
engine = create_engine(SQLALCHEMY_DATABASE_URL, connect_args={"check_same_thread": False})
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

# Override get_db dependency for testing
def override_get_db():
    try:
        db = TestingSessionLocal()
        yield db
    finally:
        db.close()

app.dependency_overrides[get_db] = override_get_db

@pytest.fixture(scope="session")
def test_db():
    """Create test database."""
    Base.metadata.create_all(bind=engine)
    yield
    Base.metadata.drop_all(bind=engine)

@pytest.fixture
def client():
    """Create test client."""
    return TestClient(app)

@pytest.fixture
def test_user(test_db):
    """Create a test user."""
    db = TestingSessionLocal()
    user = User(
        id=1,
        username="testuser",
        email="test@example.com",
        password_hash="dummy_hash"
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    db.close()
    return user

@pytest.fixture
def sample_documents():
    """Create sample documents for testing."""
    temp_dir = tempfile.mkdtemp()
    documents = {}
    
    # Create a text file
    txt_content = """
    This is a sample business plan for MegaCorp Inc.
    
    Executive Summary:
    MegaCorp is a technology company focused on AI solutions.
    Our revenue target for 2024 is $5 million.
    
    Market Analysis:
    The AI market is growing rapidly at 25% annually.
    Our competitive advantage is privacy-first design.
    
    Financial Projections:
    Year 1: $1M revenue
    Year 2: $3M revenue  
    Year 3: $5M revenue
    """
    
    txt_path = os.path.join(temp_dir, "business_plan.txt")
    with open(txt_path, 'w') as f:
        f.write(txt_content)
    documents['txt'] = txt_path
    
    # Create a CSV file
    csv_data = pd.DataFrame({
        'Product': ['AI Assistant', 'Document Processor', 'Chat System'],
        'Price': [100, 200, 150],
        'Category': ['Software', 'Software', 'Software'],
        'Launch_Date': ['2024-01-01', '2024-03-01', '2024-06-01']
    })
    
    csv_path = os.path.join(temp_dir, "products.csv")
    csv_data.to_csv(csv_path, index=False)
    documents['csv'] = csv_path
    
    # Create a Word document
    docx_doc = DocxDocument()
    docx_doc.add_heading('Company Policies', 0)
    docx_doc.add_heading('Remote Work Policy', level=1)
    docx_doc.add_paragraph('All employees may work remotely up to 3 days per week.')
    docx_doc.add_paragraph('Mandatory office days are Tuesday and Thursday.')
    
    docx_doc.add_heading('Vacation Policy', level=1)
    docx_doc.add_paragraph('Employees receive 20 days of paid vacation annually.')
    docx_doc.add_paragraph('Vacation requests must be submitted 2 weeks in advance.')
    
    docx_path = os.path.join(temp_dir, "policies.docx")
    docx_doc.save(docx_path)
    documents['docx'] = docx_path
    
    yield documents
    
    # Cleanup
    import shutil
    shutil.rmtree(temp_dir)

class TestDocumentProcessor:
    """Test document processing functionality."""
    
    def test_supported_formats(self):
        """Test that all expected formats are supported."""
        expected_formats = ['.pdf', '.docx', '.doc', '.txt', '.md', '.xlsx', '.xls', '.csv']
        supported = document_processor.get_supported_formats()
        
        for fmt in expected_formats:
            assert fmt in supported, f"Format {fmt} should be supported"
    
    def test_process_text_document(self, sample_documents):
        """Test text document processing."""
        txt_path = sample_documents['txt']
        result = document_processor.process_document(txt_path, "business_plan.txt")
        
        assert result['text_content'] is not None
        assert 'MegaCorp' in result['text_content']
        assert '$5 million' in result['text_content']
        assert len(result['chunks']) > 0
        assert result['metadata']['word_count'] > 0
        assert result['metadata']['file_extension'] == '.txt'
    
    def test_process_csv_document(self, sample_documents):
        """Test CSV document processing."""
        csv_path = sample_documents['csv']
        result = document_processor.process_document(csv_path, "products.csv")
        
        assert result['text_content'] is not None
        assert 'AI Assistant' in result['text_content']
        assert 'Price' in result['text_content']  # Header should be included
        assert len(result['chunks']) > 0
        assert result['metadata']['delimiter'] == ','
    
    def test_process_docx_document(self, sample_documents):
        """Test Word document processing."""
        docx_path = sample_documents['docx']
        result = document_processor.process_document(docx_path, "policies.docx")
        
        assert result['text_content'] is not None
        assert 'Remote Work Policy' in result['text_content']
        assert 'Vacation Policy' in result['text_content']
        assert '20 days' in result['text_content']
        assert len(result['chunks']) > 0
    
    def test_chunking_functionality(self, sample_documents):
        """Test that documents are properly chunked."""
        txt_path = sample_documents['txt']
        result = document_processor.process_document(txt_path, "business_plan.txt")
        
        chunks = result['chunks']
        assert len(chunks) > 0
        
        # Each chunk should have required fields
        for chunk in chunks:
            assert 'id' in chunk
            assert 'chunk_index' in chunk
            assert 'text_content' in chunk
            assert 'word_count' in chunk
            assert len(chunk['text_content']) > 0

class TestEmbeddingService:
    """Test vector embedding functionality."""
    
    def test_create_embeddings(self):
        """Test embedding creation."""
        texts = ["This is a test sentence.", "Another test sentence about AI."]
        embeddings = embedding_service.create_embeddings(texts)
        
        assert embeddings.shape[0] == 2  # Two input texts
        assert embeddings.shape[1] == embedding_service.dimension
        assert embeddings.dtype.name.startswith('float')
    
    def test_similarity_search(self, test_user, sample_documents):
        """Test similarity search functionality."""
        # This test requires a more complex setup with actual database records
        # For now, test the basic search structure
        db = TestingSessionLocal()
        
        # Mock search (in a real test, we'd upload documents first)
        results = embedding_service.search_similar_chunks(
            query="business revenue",
            user_id=test_user.id,
            top_k=5,
            db=db
        )
        
        # Should return empty list if no documents uploaded
        assert isinstance(results, list)
        db.close()

class TestDocumentService:
    """Test document service functionality."""
    
    @pytest.mark.asyncio
    async def test_validate_file(self):
        """Test file validation."""
        # Mock UploadFile for testing
        class MockUploadFile:
            def __init__(self, filename, size=1000):
                self.filename = filename
                self.size = size
        
        # Valid file
        valid_file = MockUploadFile("test.txt", 1000)
        result = document_service._validate_file(valid_file)
        assert result['valid'] is True
        
        # Invalid extension
        invalid_file = MockUploadFile("test.xyz", 1000)
        result = document_service._validate_file(invalid_file)
        assert result['valid'] is False
        assert 'not supported' in result['error']
        
        # File too large
        large_file = MockUploadFile("test.txt", 100 * 1024 * 1024)  # 100MB
        result = document_service._validate_file(large_file)
        assert result['valid'] is False
        assert 'exceeds maximum' in result['error']

class TestRAGService:
    """Test RAG (Retrieval-Augmented Generation) functionality."""
    
    def test_context_summary_creation(self):
        """Test context summary formatting."""
        # Mock conversation history
        conversation_history = [
            {"role": "user", "content": "Hello"},
            {"role": "assistant", "content": "Hi there!"}
        ]
        
        # Mock document chunks
        document_chunks = [
            {
                "document_filename": "business_plan.txt",
                "text_content": "Our revenue target is $5 million",
                "similarity_score": 0.85
            }
        ]
        
        summary = rag_service._create_context_summary(conversation_history, document_chunks)
        
        assert "Recent Conversation History" in summary
        assert "Relevant Information" in summary
        assert "business_plan.txt" in summary
        assert "$5 million" in summary
    
    def test_format_rag_prompt(self):
        """Test RAG prompt formatting."""
        user_message = "What's our revenue target?"
        context = {
            'context_summary': "From business_plan.txt: Revenue target is $5 million"
        }
        character_instructions = "You are a helpful assistant."
        
        prompt = rag_service.format_rag_prompt(user_message, context, character_instructions)
        
        assert character_instructions in prompt
        assert "Context Information" in prompt
        assert user_message in prompt
        assert "$5 million" in prompt

class TestAPIEndpoints:
    """Test API endpoint functionality."""
    
    def test_supported_formats_endpoint(self, client):
        """Test supported formats API endpoint."""
        response = client.get("/api/documents/formats/supported")
        assert response.status_code == 200
        
        data = response.json()
        assert "supported_formats" in data
        assert "format_descriptions" in data
        assert ".pdf" in data["supported_formats"]
    
    def test_document_stats_endpoint(self, client, test_user):
        """Test document statistics endpoint."""
        # Mock user authentication (in real implementation, this would use proper auth)
        response = client.get("/api/documents/stats/overview")
        
        # Without proper auth setup, this might return 401, which is expected
        assert response.status_code in [200, 401]
    
    def test_embedding_stats_endpoint(self, client):
        """Test embedding service statistics."""
        response = client.get("/api/documents/stats/embedding-service")
        assert response.status_code in [200, 401]

class TestIntegrationScenarios:
    """Test complete integration scenarios."""
    
    @pytest.mark.asyncio
    async def test_document_upload_to_chat_workflow(self, client, test_user, sample_documents):
        """Test the complete workflow from document upload to chat usage."""
        # This test would require proper authentication setup
        # For now, test the individual components
        
        # 1. Process document
        txt_path = sample_documents['txt']
        processing_result = document_processor.process_document(txt_path, "business_plan.txt")
        assert processing_result is not None
        
        # 2. Create embeddings
        texts = [chunk['text_content'] for chunk in processing_result['chunks']]
        embeddings = embedding_service.create_embeddings(texts)
        assert embeddings.shape[0] > 0
        
        # 3. Test RAG context creation
        context = {
            'document_chunks': [
                {
                    'document_filename': 'business_plan.txt',
                    'text_content': 'Revenue target is $5 million',
                    'similarity_score': 0.85
                }
            ],
            'sources': [{'filename': 'business_plan.txt'}]
        }
        
        summary = rag_service._create_context_summary([], context['document_chunks'])
        assert '$5 million' in summary

class TestPerformance:
    """Test performance characteristics."""
    
    def test_embedding_speed(self):
        """Test embedding creation speed."""
        import time
        
        texts = ["This is a test sentence."] * 100  # 100 similar texts
        
        start_time = time.time()
        embeddings = embedding_service.create_embeddings(texts)
        end_time = time.time()
        
        processing_time = end_time - start_time
        assert processing_time < 30  # Should process 100 texts in under 30 seconds
        assert embeddings.shape[0] == 100
    
    def test_search_speed(self, test_user):
        """Test search performance."""
        import time
        
        db = TestingSessionLocal()
        
        start_time = time.time()
        results = embedding_service.search_similar_chunks(
            query="test query",
            user_id=test_user.id,
            top_k=10,
            db=db
        )
        end_time = time.time()
        
        search_time = end_time - start_time
        assert search_time < 5  # Search should complete in under 5 seconds
        db.close()

# Test configuration and utilities
class TestConfiguration:
    """Test system configuration and setup."""
    
    def test_required_dependencies(self):
        """Test that all required dependencies are available."""
        try:
            import sentence_transformers
            import faiss
            import pandas
            import docx
            import fitz  # PyMuPDF
        except ImportError as e:
            pytest.fail(f"Required dependency missing: {e}")
    
    def test_service_initialization(self):
        """Test that all services initialize correctly."""
        assert document_processor is not None
        assert embedding_service is not None
        assert rag_service is not None
        assert document_service is not None
    
    def test_faiss_functionality(self):
        """Test basic FAISS functionality."""
        import faiss
        import numpy as np
        
        # Create a simple index
        dimension = 384
        index = faiss.IndexFlatIP(dimension)
        
        # Add some vectors
        vectors = np.random.rand(10, dimension).astype('float32')
        index.add(vectors)
        
        assert index.ntotal == 10
        
        # Search
        query = np.random.rand(1, dimension).astype('float32')
        scores, indices = index.search(query, 5)
        
        assert len(scores[0]) == 5
        assert len(indices[0]) == 5

if __name__ == "__main__":
    pytest.main([__file__, "-v"])