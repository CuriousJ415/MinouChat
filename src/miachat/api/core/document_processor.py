"""
Document processing service for extracting text from various file formats.
"""

import os
import re
import uuid
import hashlib
import logging
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import fitz  # PyMuPDF
import docx
import pandas as pd
from pathlib import Path
import json
import tiktoken

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing various document formats and extracting text."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize the document processor.
        
        Args:
            chunk_size: Maximum size of text chunks in tokens
            chunk_overlap: Number of tokens to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # Will attempt docx processing
            '.txt': self._process_text,
            '.md': self._process_text,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.csv': self._process_csv
        }
        
        # Initialize tokenizer for chunk sizing
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except Exception as e:
            logger.warning(f"Could not load tiktoken encoder: {e}")
            self.tokenizer = None
    
    def is_supported(self, filename: str) -> bool:
        """Check if a file format is supported.
        
        Args:
            filename: Name of the file
            
        Returns:
            True if format is supported, False otherwise
        """
        ext = Path(filename).suffix.lower()
        return ext in self.supported_formats
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats.
        
        Returns:
            List of supported file extensions
        """
        return list(self.supported_formats.keys())
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process a document and extract text content with metadata.
        
        Args:
            file_path: Path to the file
            filename: Original filename
            
        Returns:
            Dictionary containing extracted text, metadata, and chunks
        """
        try:
            ext = Path(filename).suffix.lower()
            
            if ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {ext}")
            
            # Get file stats
            file_size = os.path.getsize(file_path)
            
            # Process file based on type
            processor = self.supported_formats[ext]
            text_content, metadata = processor(file_path)
            
            # Calculate content hash
            content_hash = hashlib.sha256(text_content.encode('utf-8')).hexdigest()
            
            # Create text chunks
            chunks = self._create_chunks(text_content)
            
            result = {
                'text_content': text_content,
                'content_hash': content_hash,
                'file_size': file_size,
                'chunks': chunks,
                'metadata': {
                    'original_filename': filename,
                    'file_extension': ext,
                    'processing_date': datetime.utcnow().isoformat(),
                    'word_count': len(text_content.split()),
                    'character_count': len(text_content),
                    'chunk_count': len(chunks),
                    **metadata
                }
            }
            
            logger.info(f"Processed document {filename}: {len(text_content)} chars, {len(chunks)} chunks")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            raise
    
    def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = fitz.open(file_path)
            text_content = ""
            metadata = {
                'page_count': len(doc),
                'pdf_metadata': doc.metadata
            }
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                text_content += f"\n--- Page {page_num + 1} ---\n{text}\n"
            
            doc.close()
            
            # Clean up text
            text_content = self._clean_text(text_content)
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            # Try to get document properties
            try:
                core_props = doc.core_properties
                metadata.update({
                    'author': core_props.author or '',
                    'title': core_props.title or '',
                    'subject': core_props.subject or '',
                    'created': core_props.created.isoformat() if core_props.created else '',
                    'modified': core_props.modified.isoformat() if core_props.modified else ''
                })
            except Exception:
                pass
            
            text_content = self._clean_text(text_content)
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def _process_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process a plain text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text_content = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text_content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text_content:
                raise ValueError("Could not decode text file with any supported encoding")
            
            metadata = {
                'line_count': len(text_content.split('\n')),
                'encoding_used': encoding
            }
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise
    
    def _process_excel(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process an Excel file.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # Read all sheets
            excel_file = pd.read_excel(file_path, sheet_name=None)
            text_content = ""
            sheet_info = {}
            
            for sheet_name, df in excel_file.items():
                # Convert DataFrame to text representation
                sheet_text = f"\n--- Sheet: {sheet_name} ---\n"
                
                # Add column headers
                headers = " | ".join(str(col) for col in df.columns)
                sheet_text += f"Headers: {headers}\n\n"
                
                # Add data rows (limit to prevent enormous text)
                max_rows = 100  # Limit rows to prevent massive text
                for i, row in df.head(max_rows).iterrows():
                    row_text = " | ".join(str(val) for val in row.values)
                    sheet_text += f"{row_text}\n"
                
                if len(df) > max_rows:
                    sheet_text += f"... ({len(df) - max_rows} more rows)\n"
                
                text_content += sheet_text
                
                sheet_info[sheet_name] = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': list(df.columns)
                }
            
            metadata = {
                'sheet_count': len(excel_file),
                'sheets': sheet_info
            }
            
            text_content = self._clean_text(text_content)
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {e}")
            raise
    
    def _process_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process a CSV file.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # Try to detect delimiter
            with open(file_path, 'r', encoding='utf-8') as f:
                sample = f.read(1024)
            
            delimiter = ','
            if ';' in sample:
                delimiter = ';'
            elif '\t' in sample:
                delimiter = '\t'
            
            df = pd.read_csv(file_path, delimiter=delimiter)
            
            # Convert to text
            text_content = f"CSV Data (delimiter: {delimiter})\n\n"
            
            # Add headers
            headers = " | ".join(str(col) for col in df.columns)
            text_content += f"Headers: {headers}\n\n"
            
            # Add data (limit rows)
            max_rows = 100
            for i, row in df.head(max_rows).iterrows():
                row_text = " | ".join(str(val) for val in row.values)
                text_content += f"{row_text}\n"
            
            if len(df) > max_rows:
                text_content += f"... ({len(df) - max_rows} more rows)\n"
            
            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'delimiter': delimiter
            }
            
            text_content = self._clean_text(text_content)
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove control characters but keep newlines and tabs
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Trim
        text = text.strip()
        
        return text
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks for better retrieval.
        
        Args:
            text: Text to chunk
            
        Returns:
            List of chunk dictionaries
        """
        chunks = []
        
        if not text.strip():
            return chunks
        
        # If tokenizer is available, use token-based chunking
        if self.tokenizer:
            chunks = self._create_token_chunks(text)
        else:
            # Fallback to character-based chunking
            chunks = self._create_char_chunks(text)
        
        return chunks
    
    def _create_token_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create chunks based on token count.
        
        Args:
            text: Text to chunk
            
        Returns:
            List of chunk dictionaries
        """
        chunks = []
        
        try:
            # Tokenize the text
            tokens = self.tokenizer.encode(text)
            
            # Create overlapping chunks
            start = 0
            chunk_index = 0
            
            while start < len(tokens):
                end = min(start + self.chunk_size, len(tokens))
                chunk_tokens = tokens[start:end]
                chunk_text = self.tokenizer.decode(chunk_tokens)
                
                # Clean up chunk text
                chunk_text = chunk_text.strip()
                
                if chunk_text:
                    chunk = {
                        'id': str(uuid.uuid4()),
                        'chunk_index': chunk_index,
                        'text_content': chunk_text,
                        'start_char': None,  # Could calculate if needed
                        'end_char': None,
                        'word_count': len(chunk_text.split()),
                        'token_count': len(chunk_tokens),
                        'chunk_type': 'paragraph'
                    }
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Move start position with overlap
                if end == len(tokens):
                    break
                start = max(start + self.chunk_size - self.chunk_overlap, start + 1)
            
        except Exception as e:
            logger.error(f"Error creating token chunks: {e}")
            # Fallback to character chunking
            return self._create_char_chunks(text)
        
        return chunks
    
    def _create_char_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create chunks based on character count (fallback method).
        
        Args:
            text: Text to chunk
            
        Returns:
            List of chunk dictionaries
        """
        chunks = []
        char_chunk_size = self.chunk_size * 4  # Approximate 4 chars per token
        char_overlap = self.chunk_overlap * 4
        
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = min(start + char_chunk_size, len(text))
            
            # Try to end at a sentence or paragraph boundary
            if end < len(text):
                # Look for sentence ending
                for i in range(end, max(start + char_chunk_size // 2, start + 1), -1):
                    if text[i-1] in '.!?\n':
                        end = i
                        break
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk = {
                    'id': str(uuid.uuid4()),
                    'chunk_index': chunk_index,
                    'text_content': chunk_text,
                    'start_char': start,
                    'end_char': end,
                    'word_count': len(chunk_text.split()),
                    'token_count': None,
                    'chunk_type': 'paragraph'
                }
                chunks.append(chunk)
                chunk_index += 1
            
            if end == len(text):
                break
            start = max(end - char_overlap, start + 1)
        
        return chunks

# Global document processor instance
document_processor = DocumentProcessor()