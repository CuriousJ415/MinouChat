"""
Export Service for MinouChat

Handles exporting conversations and documents to various formats:
- PDF (.pdf) - Final documents, sharing
- Word (.docx) - Formal reports, editable documents
- Markdown (.md) - Technical docs, GitHub-friendly
- Plain text (.txt) - Universal, simple notes

Supports LLM-powered document generation with category-specific templates:
- Assistant: Research Summary, Decision Log, Project Brief, Email Draft
- Coach: Goal Progress, Weekly Review, Action Plan, Values Assessment, Life Areas
- Teacher: Lesson Summary, Study Guide, Concept Explanation, Practice Problems
- Friend: Advice & Wisdom, Collaborative Short Story
- Creative: Story Summary (brief/detailed)
"""

import os
import io
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional, Union, Tuple
from enum import Enum
from dataclasses import dataclass

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

# Word document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class ExportFormat(Enum):
    """Supported export formats."""
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "md"
    TEXT = "txt"


class DocumentType(Enum):
    """Types of documents that can be generated."""
    # Assistant documents
    RESEARCH_SUMMARY = "research_summary"
    DECISION_LOG = "decision_log"
    PROJECT_BRIEF = "project_brief"
    EMAIL_DRAFT = "email_draft"

    # Coach documents
    GOAL_PROGRESS_REPORT = "goal_progress_report"
    WEEKLY_REVIEW = "weekly_review"
    MONTHLY_REVIEW = "monthly_review"
    ACTION_PLAN = "action_plan"
    ACCOUNTABILITY_SUMMARY = "accountability_summary"
    REFLECTION_JOURNAL = "reflection_journal"
    HABIT_TRACKER_SUMMARY = "habit_tracker_summary"
    VALUES_ASSESSMENT = "values_assessment"
    LIFE_AREAS_ASSESSMENT = "life_areas_assessment"
    # New Coach assessment documents
    ONBOARDING_REPORT = "onboarding_report"
    STRENGTHS_PROFILE = "strengths_profile"
    WHEEL_OF_LIFE = "wheel_of_life"

    # Teacher documents
    LESSON_SUMMARY = "lesson_summary"
    STUDY_GUIDE = "study_guide"
    CONCEPT_EXPLANATION = "concept_explanation"
    PRACTICE_PROBLEMS = "practice_problems"
    LEARNING_PROGRESS = "learning_progress"
    QA_DOCUMENT = "qa_document"

    # Friend documents
    ADVICE_WISDOM = "advice_wisdom"
    COLLABORATIVE_STORY = "collaborative_story"

    # Creative/Roleplay documents
    STORY_SUMMARY_BRIEF = "story_summary_brief"
    STORY_SUMMARY_DETAILED = "story_summary_detailed"


@dataclass
class DocumentTypeInfo:
    """Metadata about a document type."""
    doc_type: DocumentType
    name: str
    description: str
    category: str
    tone: str
    template_style: str  # "hybrid", "freeform", "bullet_list", "prose"


# Document type registry with metadata
DOCUMENT_TYPES: Dict[DocumentType, DocumentTypeInfo] = {
    # Assistant
    DocumentType.RESEARCH_SUMMARY: DocumentTypeInfo(
        doc_type=DocumentType.RESEARCH_SUMMARY,
        name="Research Summary",
        description="Organized findings with key insights and recommendations",
        category="Assistant",
        tone="professional, direct",
        template_style="hybrid"
    ),
    DocumentType.DECISION_LOG: DocumentTypeInfo(
        doc_type=DocumentType.DECISION_LOG,
        name="Decision Log",
        description="Decisions made, rationale, and alternatives considered",
        category="Assistant",
        tone="professional, direct",
        template_style="hybrid"
    ),
    DocumentType.PROJECT_BRIEF: DocumentTypeInfo(
        doc_type=DocumentType.PROJECT_BRIEF,
        name="Project Brief",
        description="Goals, scope, timeline, and stakeholders",
        category="Assistant",
        tone="professional, direct",
        template_style="hybrid"
    ),
    DocumentType.EMAIL_DRAFT: DocumentTypeInfo(
        doc_type=DocumentType.EMAIL_DRAFT,
        name="Email Draft",
        description="Professional email based on conversation points",
        category="Assistant",
        tone="professional, direct",
        template_style="freeform"
    ),

    # Coach
    DocumentType.GOAL_PROGRESS_REPORT: DocumentTypeInfo(
        doc_type=DocumentType.GOAL_PROGRESS_REPORT,
        name="Goal Progress Report",
        description="Current goals, progress made, obstacles, next steps",
        category="Coach",
        tone="encouraging, direct, accountability-focused",
        template_style="hybrid"
    ),
    DocumentType.WEEKLY_REVIEW: DocumentTypeInfo(
        doc_type=DocumentType.WEEKLY_REVIEW,
        name="Weekly Review",
        description="Wins, challenges, lessons learned, focus areas",
        category="Coach",
        tone="encouraging, direct, accountability-focused",
        template_style="hybrid"
    ),
    DocumentType.MONTHLY_REVIEW: DocumentTypeInfo(
        doc_type=DocumentType.MONTHLY_REVIEW,
        name="Monthly Review",
        description="Month's achievements, patterns, growth areas",
        category="Coach",
        tone="encouraging, direct, accountability-focused",
        template_style="hybrid"
    ),
    DocumentType.ACTION_PLAN: DocumentTypeInfo(
        doc_type=DocumentType.ACTION_PLAN,
        name="Action Plan",
        description="Step-by-step plan with milestones and deadlines",
        category="Coach",
        tone="encouraging, direct, accountability-focused",
        template_style="hybrid"
    ),
    DocumentType.ACCOUNTABILITY_SUMMARY: DocumentTypeInfo(
        doc_type=DocumentType.ACCOUNTABILITY_SUMMARY,
        name="Accountability Summary",
        description="Commitments made, completion status, patterns",
        category="Coach",
        tone="encouraging, direct, accountability-focused",
        template_style="hybrid"
    ),
    DocumentType.REFLECTION_JOURNAL: DocumentTypeInfo(
        doc_type=DocumentType.REFLECTION_JOURNAL,
        name="Reflection Journal",
        description="Insights, growth areas, mindset shifts",
        category="Coach",
        tone="encouraging, direct, accountability-focused",
        template_style="hybrid"
    ),
    DocumentType.HABIT_TRACKER_SUMMARY: DocumentTypeInfo(
        doc_type=DocumentType.HABIT_TRACKER_SUMMARY,
        name="Habit Tracker Summary",
        description="Habits discussed, streaks, recommendations",
        category="Coach",
        tone="encouraging, direct, accountability-focused",
        template_style="hybrid"
    ),
    DocumentType.VALUES_ASSESSMENT: DocumentTypeInfo(
        doc_type=DocumentType.VALUES_ASSESSMENT,
        name="Values & Principles Assessment",
        description="Core values, guiding principles, identity statements",
        category="Coach",
        tone="encouraging, direct, accountability-focused",
        template_style="hybrid"
    ),
    DocumentType.LIFE_AREAS_ASSESSMENT: DocumentTypeInfo(
        doc_type=DocumentType.LIFE_AREAS_ASSESSMENT,
        name="Life Areas Assessment",
        description="Assessment across work, health, relationships, growth",
        category="Coach",
        tone="encouraging, direct, accountability-focused",
        template_style="hybrid"
    ),
    DocumentType.ONBOARDING_REPORT: DocumentTypeInfo(
        doc_type=DocumentType.ONBOARDING_REPORT,
        name="Onboarding Report",
        description="Complete initial assessment with values, strengths, goals, and recommendations",
        category="Coach",
        tone="encouraging, insightful, comprehensive",
        template_style="hybrid"
    ),
    DocumentType.STRENGTHS_PROFILE: DocumentTypeInfo(
        doc_type=DocumentType.STRENGTHS_PROFILE,
        name="Strengths Profile",
        description="Character strengths with application suggestions and development tips",
        category="Coach",
        tone="affirming, practical, empowering",
        template_style="hybrid"
    ),
    DocumentType.WHEEL_OF_LIFE: DocumentTypeInfo(
        doc_type=DocumentType.WHEEL_OF_LIFE,
        name="Wheel of Life Snapshot",
        description="Life satisfaction scores across 10 domains with insights",
        category="Coach",
        tone="balanced, insightful, action-oriented",
        template_style="hybrid"
    ),

    # Teacher
    DocumentType.LESSON_SUMMARY: DocumentTypeInfo(
        doc_type=DocumentType.LESSON_SUMMARY,
        name="Lesson Summary",
        description="Key concepts taught, examples, important points",
        category="Teacher",
        tone="academic, approachable, adaptable",
        template_style="hybrid"
    ),
    DocumentType.STUDY_GUIDE: DocumentTypeInfo(
        doc_type=DocumentType.STUDY_GUIDE,
        name="Study Guide",
        description="Topics formatted for review and studying",
        category="Teacher",
        tone="academic, approachable, adaptable",
        template_style="hybrid"
    ),
    DocumentType.CONCEPT_EXPLANATION: DocumentTypeInfo(
        doc_type=DocumentType.CONCEPT_EXPLANATION,
        name="Concept Explanation",
        description="Deep-dive on a specific topic with examples",
        category="Teacher",
        tone="academic, approachable, adaptable",
        template_style="hybrid"
    ),
    DocumentType.PRACTICE_PROBLEMS: DocumentTypeInfo(
        doc_type=DocumentType.PRACTICE_PROBLEMS,
        name="Practice Problems",
        description="Exercises based on material covered",
        category="Teacher",
        tone="academic, approachable, adaptable",
        template_style="hybrid"
    ),
    DocumentType.LEARNING_PROGRESS: DocumentTypeInfo(
        doc_type=DocumentType.LEARNING_PROGRESS,
        name="Learning Progress",
        description="Skills developed, areas to improve, next topics",
        category="Teacher",
        tone="academic, approachable, adaptable",
        template_style="hybrid"
    ),
    DocumentType.QA_DOCUMENT: DocumentTypeInfo(
        doc_type=DocumentType.QA_DOCUMENT,
        name="Q&A Document",
        description="Questions asked and answers in reference format",
        category="Teacher",
        tone="academic, approachable, adaptable",
        template_style="hybrid"
    ),

    # Friend
    DocumentType.ADVICE_WISDOM: DocumentTypeInfo(
        doc_type=DocumentType.ADVICE_WISDOM,
        name="Advice & Wisdom",
        description="Key advice and perspectives shared",
        category="Friend",
        tone="factual",
        template_style="bullet_list"
    ),
    DocumentType.COLLABORATIVE_STORY: DocumentTypeInfo(
        doc_type=DocumentType.COLLABORATIVE_STORY,
        name="Collaborative Short Story",
        description="Story created together, appends new content",
        category="Friend",
        tone="flexible, narrative",
        template_style="prose"
    ),

    # Creative/Roleplay
    DocumentType.STORY_SUMMARY_BRIEF: DocumentTypeInfo(
        doc_type=DocumentType.STORY_SUMMARY_BRIEF,
        name="Story Summary (Brief)",
        description="Concise overview of the story so far",
        category="Creative",
        tone="novel prose, third person",
        template_style="prose"
    ),
    DocumentType.STORY_SUMMARY_DETAILED: DocumentTypeInfo(
        doc_type=DocumentType.STORY_SUMMARY_DETAILED,
        name="Story Summary (Detailed)",
        description="Comprehensive narrative of the story",
        category="Creative",
        tone="novel prose, third person",
        template_style="prose"
    ),
}

# Category to document types mapping
CATEGORY_DOCUMENTS: Dict[str, List[DocumentType]] = {
    "Assistant": [
        DocumentType.RESEARCH_SUMMARY,
        DocumentType.DECISION_LOG,
        DocumentType.PROJECT_BRIEF,
        DocumentType.EMAIL_DRAFT,
    ],
    "Coach": [
        DocumentType.GOAL_PROGRESS_REPORT,
        DocumentType.WEEKLY_REVIEW,
        DocumentType.MONTHLY_REVIEW,
        DocumentType.ACTION_PLAN,
        DocumentType.ACCOUNTABILITY_SUMMARY,
        DocumentType.REFLECTION_JOURNAL,
        DocumentType.HABIT_TRACKER_SUMMARY,
        DocumentType.VALUES_ASSESSMENT,
        DocumentType.LIFE_AREAS_ASSESSMENT,
        DocumentType.ONBOARDING_REPORT,
        DocumentType.STRENGTHS_PROFILE,
        DocumentType.WHEEL_OF_LIFE,
    ],
    "Teacher": [
        DocumentType.LESSON_SUMMARY,
        DocumentType.STUDY_GUIDE,
        DocumentType.CONCEPT_EXPLANATION,
        DocumentType.PRACTICE_PROBLEMS,
        DocumentType.LEARNING_PROGRESS,
        DocumentType.QA_DOCUMENT,
    ],
    "Friend": [
        DocumentType.ADVICE_WISDOM,
        DocumentType.COLLABORATIVE_STORY,
    ],
    "Creative": [
        DocumentType.STORY_SUMMARY_BRIEF,
        DocumentType.STORY_SUMMARY_DETAILED,
    ],
}

# Life Areas for Coach assessments
LIFE_AREAS = [
    "Work",
    "Environment (home, office, city, aesthetics)",
    "Romance/Relationship",
    "Family",
    "Community",
    "Friendships",
    "Health (sleep, alcohol, exercise, diet, weight, medical)",
    "Spiritual",
    "Finances",
    "Creativity",
    "Personal Growth",
]

# LLM Prompt Templates for each document type
DOCUMENT_PROMPTS: Dict[DocumentType, str] = {
    # =========================================================================
    # ASSISTANT DOCUMENTS
    # =========================================================================
    DocumentType.RESEARCH_SUMMARY: """You are generating a professional Research Summary document.

TONE: Professional and direct
FORMAT: Hybrid template with required and optional sections

REQUIRED SECTIONS:
## Executive Summary
(2-3 sentences summarizing the key findings)

## Key Findings
(Numbered list of main discoveries/insights from the conversation)

## Recommendations
(Actionable next steps based on the findings)

OPTIONAL SECTIONS (include only if relevant to the conversation):
## Sources & References
## Open Questions
## Methodology
## Data Points

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document content in markdown format. Be professional and direct.""",

    DocumentType.DECISION_LOG: """You are generating a professional Decision Log document.

TONE: Professional and direct
FORMAT: Hybrid template with required and optional sections

REQUIRED SECTIONS:
## Decisions Made
(List each decision with date/context if available)

## Rationale
(Why each decision was made)

## Alternatives Considered
(Other options that were discussed)

OPTIONAL SECTIONS (include only if relevant):
## Impact Assessment
## Action Items
## Follow-up Required

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document content in markdown format. Be professional and direct.""",

    DocumentType.PROJECT_BRIEF: """You are generating a professional Project Brief document.

TONE: Professional and direct
FORMAT: Hybrid template with required and optional sections

REQUIRED SECTIONS:
## Project Overview
(Brief description of the project)

## Goals & Objectives
(What the project aims to achieve)

## Scope
(What's included and excluded)

OPTIONAL SECTIONS (include only if discussed):
## Timeline & Milestones
## Stakeholders
## Resources Required
## Success Metrics
## Risks & Mitigation

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document content in markdown format. Be professional and direct.""",

    DocumentType.EMAIL_DRAFT: """You are drafting a professional email based on the conversation.

TONE: Professional and direct
FORMAT: Standard email format

Generate an email with:
- Subject line
- Greeting
- Body (clear, concise paragraphs)
- Call to action or next steps
- Professional closing

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the email in a ready-to-send format.""",

    # =========================================================================
    # COACH DOCUMENTS
    # =========================================================================
    DocumentType.GOAL_PROGRESS_REPORT: """You are a coach generating a Goal Progress Report for your client.

TONE: Encouraging, direct, and accountability-focused (based on Atomic Habits principles)
PERSPECTIVE: Write as the coach speaking to the client ("I noticed...", "You've made progress on...")

REQUIRED SECTIONS:
## Progress Overview
(Summary of where they stand on their goals)

## Goals & Current Status
(Each goal with progress indicators)

## Wins to Celebrate
(Acknowledge their successes, no matter how small)

## Obstacles Identified
(Challenges they're facing)

## Next Steps
(Specific, actionable items)

OPTIONAL SECTIONS:
## Identity Reinforcement
(Connect progress to who they're becoming)
## Habit Stacking Opportunities
## Environment Design Suggestions

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format. Be encouraging but hold them accountable.""",

    DocumentType.WEEKLY_REVIEW: """You are a coach generating a Weekly Review for your client.

TONE: Encouraging, direct, and accountability-focused
PERSPECTIVE: Write as the coach speaking to the client

REQUIRED SECTIONS:
## This Week's Wins
(What went well, celebrate the progress)

## Challenges Faced
(What was difficult, what got in the way)

## Lessons Learned
(Insights and realizations from this week)

## Focus for Next Week
(Key priorities and intentions)

OPTIONAL SECTIONS:
## Habit Check-in
## Energy & Motivation Level
## Gratitude Moments

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.MONTHLY_REVIEW: """You are a coach generating a Monthly Review for your client.

TONE: Encouraging, direct, and accountability-focused
PERSPECTIVE: Write as the coach speaking to the client

REQUIRED SECTIONS:
## Month at a Glance
(Overview of the month's journey)

## Major Achievements
(Significant wins and milestones)

## Growth Areas
(Where they've developed)

## Patterns Noticed
(Recurring themes, habits, or behaviors)

## Focus for Next Month
(Priorities and goals)

OPTIONAL SECTIONS:
## Life Areas Check-in
## Values Alignment
## Momentum Assessment

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.ACTION_PLAN: """You are a coach generating an Action Plan for your client.

TONE: Encouraging, direct, and accountability-focused
PERSPECTIVE: Write as the coach speaking to the client

REQUIRED SECTIONS:
## Goal Statement
(Clear articulation of what they're working toward)

## Action Steps
(Numbered, specific steps with timeframes)

## Milestones
(Key checkpoints along the way)

## Potential Obstacles
(What might get in the way)

## Support & Resources
(What they need to succeed)

OPTIONAL SECTIONS:
## Habit Cues & Rewards
## Accountability Measures
## Contingency Plans

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.ACCOUNTABILITY_SUMMARY: """You are a coach generating an Accountability Summary.

TONE: Encouraging, direct, and accountability-focused
PERSPECTIVE: Write as the coach speaking to the client

REQUIRED SECTIONS:
## Commitments Made
(What they said they would do)

## Completion Status
(What was done vs. not done)

## Patterns Observed
(What the coach notices about their behavior)

## Reflection Questions
(Questions for the client to consider)

OPTIONAL SECTIONS:
## Identity Check
## Momentum Assessment
## Adjustment Recommendations

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format. Be honest but supportive.""",

    DocumentType.REFLECTION_JOURNAL: """You are a coach generating a Reflection Journal entry.

TONE: Encouraging, direct, thoughtful
PERSPECTIVE: Write as the coach capturing the client's insights

REQUIRED SECTIONS:
## Key Insights
(Major realizations from the conversation)

## Growth Moments
(Where they showed growth or awareness)

## Mindset Shifts
(Changes in perspective or thinking)

## Questions to Sit With
(Deeper questions for continued reflection)

OPTIONAL SECTIONS:
## Emotional Themes
## Values Connections
## Future Aspirations

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.HABIT_TRACKER_SUMMARY: """You are a coach generating a Habit Tracker Summary.

TONE: Encouraging, direct, and accountability-focused
PERSPECTIVE: Write as the coach speaking to the client

REQUIRED SECTIONS:
## Habits Being Tracked
(List of habits discussed)

## Progress & Streaks
(How they're doing with each habit)

## What's Working
(Successful strategies and cues)

## What Needs Adjustment
(Habits struggling and why)

## Recommendations
(Specific suggestions based on Atomic Habits principles)

OPTIONAL SECTIONS:
## Habit Stacking Opportunities
## Environment Design
## Identity Reinforcement

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.VALUES_ASSESSMENT: """You are a coach generating a Values & Principles Assessment.

TONE: Encouraging, direct, reflective
PERSPECTIVE: Write as the coach capturing the client's core values

REQUIRED SECTIONS:
## Core Values Identified
(List the values that emerged from conversation)

## Guiding Principles
(Rules or principles they want to live by)

## Identity Statements
("I am someone who...")

## Values in Action
(How these values show up in their life)

OPTIONAL SECTIONS:
## Values Conflicts
## Growth Edge
## Legacy Vision

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.LIFE_AREAS_ASSESSMENT: """You are a coach generating a Life Areas Assessment.

TONE: Encouraging, direct, holistic
PERSPECTIVE: Write as the coach assessing the client's life balance

Assess each of these life areas based on the conversation:
- Work
- Environment (home, office, city, aesthetics)
- Romance/Relationship
- Family
- Community
- Friendships
- Health (sleep, alcohol, exercise, diet, weight, medical)
- Spiritual
- Finances
- Creativity
- Personal Growth

REQUIRED SECTIONS:
## Life Areas Overview
(Brief assessment of each area mentioned)

## Areas of Strength
(Where they're thriving)

## Areas for Growth
(Where attention is needed)

## Interconnections
(How areas affect each other)

## Priorities Identified
(Top 2-3 areas to focus on)

OPTIONAL SECTIONS:
## Quick Wins
## Long-term Vision
## Support Needed

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.ONBOARDING_REPORT: """You are a premium life coach generating a comprehensive Onboarding Report.

TONE: Welcoming, insightful, encouraging yet honest
PERSPECTIVE: Write as the coach summarizing initial assessment findings

This is the complete initial assessment documenting who the client is and their starting point.
Extract from the conversation: values, strengths, life satisfaction scores, goals, and patterns.

REQUIRED SECTIONS:
## Welcome & Introduction
(Personalized acknowledgment of their commitment to growth)

## Values Discovery
(Their core values identified from conversation - aim for 5-7)
- List each value with brief explanation of how it showed up

## Character Strengths Profile
(Top 5 strengths identified using VIA framework)
- Wisdom: Creativity, Curiosity, Judgment, Love of Learning, Perspective
- Courage: Bravery, Perseverance, Honesty, Zest
- Humanity: Love, Kindness, Social Intelligence
- Justice: Teamwork, Fairness, Leadership
- Temperance: Forgiveness, Humility, Prudence, Self-Regulation
- Transcendence: Gratitude, Hope, Humor, Spirituality

## Life Satisfaction Snapshot
(Rate each area 1-10 if discussed, otherwise note "to be assessed")
- Career/Work, Finances, Health, Relationships, Family
- Friendships, Personal Growth, Fun/Recreation, Environment, Contribution

## Primary Goals
(What they want to achieve - use SMART format)

## Patterns & Saboteurs Noticed
(Any limiting beliefs or self-sabotaging patterns identified)
- The Judge, Controller, Pleaser, Victim, Avoider, Hyper-Achiever, etc.

## Recommended Focus Areas
(Top 2-3 priorities for initial coaching work)

OPTIONAL SECTIONS:
## Quick Wins Available
## Support Systems
## Coaching Approach Recommendations

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.STRENGTHS_PROFILE: """You are a coach generating a Character Strengths Profile.

TONE: Affirming, practical, empowering
PERSPECTIVE: Write as the coach celebrating and explaining their unique strengths

Use the VIA Character Strengths framework (24 strengths grouped by virtue).
Identify their top 5 signature strengths from the conversation.

VIA STRENGTHS REFERENCE:
- Wisdom: Creativity, Curiosity, Judgment, Love of Learning, Perspective
- Courage: Bravery, Perseverance, Honesty, Zest
- Humanity: Love, Kindness, Social Intelligence
- Justice: Teamwork, Fairness, Leadership
- Temperance: Forgiveness, Humility, Prudence, Self-Regulation
- Transcendence: Gratitude, Hope, Humor, Spirituality, Appreciation of Beauty

REQUIRED SECTIONS:
## Your Signature Strengths
(Top 5 strengths with evidence from conversation)

For each strength include:
### [Strength Name]
**What it means:** Brief definition
**How you show it:** Evidence from their stories/conversation
**How to apply it:** Practical ways to leverage this strength

## Strength Combinations
(How their strengths work together powerfully)

## Growth Edges
(Strengths that could be developed further)

## Strength-Based Action Plan
(Specific ways to use top strengths toward their goals)

OPTIONAL SECTIONS:
## Strengths to Watch
(When overused, these strengths can become liabilities)
## Strengths Affirmations
(Identity statements: "I am someone who...")

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.WHEEL_OF_LIFE: """You are a coach generating a Wheel of Life assessment snapshot.

TONE: Balanced, insightful, action-oriented
PERSPECTIVE: Write as the coach providing a holistic life assessment

The Wheel of Life rates satisfaction (1-10) across 10 life domains.
Extract scores from conversation if discussed, or note patterns.

LIFE DOMAINS:
1. Career/Work - Professional fulfillment, purpose, advancement
2. Finances - Income, savings, security, relationship with money
3. Health - Physical fitness, energy, sleep, nutrition, medical
4. Relationships - Romantic partnership, intimacy, connection
5. Family - Family bonds, responsibilities, harmony
6. Friendships - Social connections, community, belonging
7. Personal Growth - Learning, self-improvement, spirituality
8. Fun & Recreation - Hobbies, leisure, play, creativity
9. Environment - Home, workspace, physical surroundings
10. Contribution - Giving back, impact, legacy, service

REQUIRED SECTIONS:
## Life Satisfaction Scores
| Domain | Score (1-10) | Notes |
|--------|--------------|-------|
(Create table with all 10 domains)

## Balance Overview
(Overall assessment - where are they thriving vs struggling?)

## Highest Scoring Areas
(What's working well - celebrate these)

## Priority Focus Areas
(Lowest 2-3 scores that need attention)

## The Ripple Effect
(How improving one area could positively impact others)

## Next Steps
(Specific actions for top priority areas)

OPTIONAL SECTIONS:
## Historical Context
(If they mentioned how things have changed)
## 90-Day Vision
(What improved scores would look like)

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    # =========================================================================
    # TEACHER DOCUMENTS
    # =========================================================================
    DocumentType.LESSON_SUMMARY: """You are a teacher generating a Lesson Summary.

TONE: Academic yet approachable, adaptable to skill level
FORMAT: Clear educational structure

REQUIRED SECTIONS:
## Lesson Overview
(What was covered)

## Key Concepts
(Main ideas explained clearly)

## Examples & Illustrations
(Concrete examples from the lesson)

## Important Points to Remember
(Takeaways and key facts)

OPTIONAL SECTIONS:
## Code Snippets (for programming topics)
## Diagrams Described (for visual concepts)
## Common Mistakes to Avoid
## Further Reading

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format. Adapt language to the learner's level shown in the conversation.""",

    DocumentType.STUDY_GUIDE: """You are a teacher generating a Study Guide.

TONE: Academic yet approachable
FORMAT: Optimized for review and memorization

REQUIRED SECTIONS:
## Topics Covered
(Overview of subjects in the guide)

## Key Terms & Definitions
(Important vocabulary and concepts)

## Main Ideas
(Core concepts to understand)

## Review Questions
(Self-test questions)

OPTIONAL SECTIONS:
## Memory Aids
## Practice Problems
## Summary Tables
## Connections Between Topics

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format. Make it easy to study from.""",

    DocumentType.CONCEPT_EXPLANATION: """You are a teacher generating a Concept Explanation.

TONE: Academic yet approachable, thorough
FORMAT: Deep-dive educational content

REQUIRED SECTIONS:
## Concept Overview
(What is this concept?)

## Detailed Explanation
(In-depth breakdown)

## Examples
(Multiple examples showing the concept)

## Common Misconceptions
(What people often get wrong)

OPTIONAL SECTIONS:
## Code Implementation (for programming)
## Visual Description (for visual concepts)
## Related Concepts
## Real-World Applications

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format. Be thorough but clear.""",

    DocumentType.PRACTICE_PROBLEMS: """You are a teacher generating Practice Problems.

TONE: Academic, encouraging
FORMAT: Exercise set with varying difficulty

REQUIRED SECTIONS:
## Overview
(What these problems cover)

## Problems
(Numbered problems with clear instructions)

## Hints (optional per problem)
(Guidance without giving away the answer)

## Solutions
(Complete solutions explained)

OPTIONAL SECTIONS:
## Difficulty Ratings
## Common Mistakes
## Extension Challenges

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format. Include a mix of difficulty levels.""",

    DocumentType.LEARNING_PROGRESS: """You are a teacher generating a Learning Progress report.

TONE: Academic, encouraging, constructive
FORMAT: Progress assessment

REQUIRED SECTIONS:
## Skills Developed
(What the learner has mastered)

## Current Understanding
(Where they are now)

## Areas to Improve
(What needs more work)

## Recommended Next Steps
(What to study next)

OPTIONAL SECTIONS:
## Learning Style Observations
## Suggested Resources
## Goals for Next Session

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format.""",

    DocumentType.QA_DOCUMENT: """You are a teacher generating a Q&A Reference Document.

TONE: Academic, clear, reference-style
FORMAT: Question and answer pairs

REQUIRED SECTIONS:
## Questions & Answers
(Each question followed by its complete answer)

OPTIONAL SECTIONS:
## Follow-up Questions
## Related Topics
## Quick Reference Summary

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the document in markdown format. Format as a clear Q&A reference.""",

    # =========================================================================
    # FRIEND DOCUMENTS
    # =========================================================================
    DocumentType.ADVICE_WISDOM: """Extract the key advice and wisdom shared in this conversation.

TONE: Factual, straightforward
FORMAT: Bullet point list

Generate a bullet-point list of:
- Key advice given
- Perspectives shared
- Wisdom or insights offered
- Helpful suggestions made

Keep each point concise and actionable.

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate a clean bullet-point list only, no additional sections.""",

    DocumentType.COLLABORATIVE_STORY: """Continue or summarize the collaborative story from this conversation.

TONE: Flexible, narrative, matching the story's established tone
FORMAT: Novel prose

Write in a narrative style that:
- Captures the story events from the conversation
- Maintains the established voice and tone
- Flows naturally as prose fiction
- Can be appended to existing story content

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate the story content in prose format. This will be appended to any existing story content.""",

    # =========================================================================
    # CREATIVE/ROLEPLAY DOCUMENTS
    # =========================================================================
    DocumentType.STORY_SUMMARY_BRIEF: """Generate a brief summary of the roleplay/story.

TONE: Novel prose, third person narrative
FORMAT: Concise narrative summary

Write a brief (1-2 paragraphs) summary that:
- Captures the key events and plot points
- Describes character interactions
- Notes any significant developments
- Written in third person, past tense

{sfw_instruction}

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate a concise story summary.""",

    DocumentType.STORY_SUMMARY_DETAILED: """Generate a detailed summary of the roleplay/story.

TONE: Novel prose, third person narrative
FORMAT: Comprehensive narrative summary

Write a detailed summary that:
- Covers all significant events in order
- Describes character development
- Captures important dialogue moments
- Notes relationship dynamics
- Includes setting and atmosphere details
- Written in third person, past tense

{sfw_instruction}

CONVERSATION CONTEXT:
{conversation}

CHARACTER: {character_name}
{custom_instructions}

Generate a comprehensive story summary.""",
}


def get_document_types_for_category(category: str) -> List[Dict[str, Any]]:
    """Get available document types for a persona category.

    Args:
        category: The persona category (Assistant, Coach, Teacher, Friend, Creative)

    Returns:
        List of document type info dictionaries
    """
    doc_types = CATEGORY_DOCUMENTS.get(category, [])
    return [
        {
            "id": dt.value,
            "name": DOCUMENT_TYPES[dt].name,
            "description": DOCUMENT_TYPES[dt].description,
        }
        for dt in doc_types
    ]


class ExportService:
    """Service for exporting conversations and content to various formats."""

    def __init__(self, output_dir: str = "./output_documents"):
        """Initialize the export service.

        Args:
            output_dir: Directory to store exported files
        """
        self.output_dir = output_dir
        self._ensure_output_directory()

    def _ensure_output_directory(self):
        """Ensure the output directory exists."""
        os.makedirs(self.output_dir, exist_ok=True)

    def export_conversation(
        self,
        messages: List[Dict[str, Any]],
        title: str,
        format: ExportFormat,
        character_name: str = "AI Assistant",
        include_system: bool = False,
        include_timestamps: bool = True
    ) -> io.BytesIO:
        """Export a conversation to the specified format.

        Args:
            messages: List of message dictionaries with 'role', 'content', 'timestamp'
            title: Title for the export
            format: Export format (pdf, docx, md, txt)
            character_name: Name of the AI character
            include_system: Whether to include system messages
            include_timestamps: Whether to include timestamps

        Returns:
            BytesIO buffer containing the exported file
        """
        # Filter out system messages if requested
        if not include_system:
            messages = [m for m in messages if m.get('role') != 'system']

        if format == ExportFormat.PDF:
            return self._export_conversation_pdf(messages, title, character_name, include_timestamps)
        elif format == ExportFormat.DOCX:
            return self._export_conversation_docx(messages, title, character_name, include_timestamps)
        elif format == ExportFormat.MARKDOWN:
            return self._export_conversation_markdown(messages, title, character_name, include_timestamps)
        elif format == ExportFormat.TEXT:
            return self._export_conversation_text(messages, title, character_name, include_timestamps)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def export_summary(
        self,
        summary_text: str,
        title: str,
        format: ExportFormat,
        character_name: str = "AI Assistant",
        key_points: Optional[List[str]] = None,
        original_message_count: int = 0
    ) -> io.BytesIO:
        """Export a conversation summary to the specified format.

        Args:
            summary_text: The LLM-generated summary
            title: Title for the export
            format: Export format
            character_name: Name of the AI character
            key_points: Optional list of key points
            original_message_count: Number of messages in original conversation

        Returns:
            BytesIO buffer containing the exported file
        """
        if format == ExportFormat.PDF:
            return self._export_summary_pdf(summary_text, title, character_name, key_points, original_message_count)
        elif format == ExportFormat.DOCX:
            return self._export_summary_docx(summary_text, title, character_name, key_points, original_message_count)
        elif format == ExportFormat.MARKDOWN:
            return self._export_summary_markdown(summary_text, title, character_name, key_points, original_message_count)
        elif format == ExportFormat.TEXT:
            return self._export_summary_text(summary_text, title, character_name, key_points, original_message_count)
        else:
            raise ValueError(f"Unsupported format: {format}")

    # =========================================================================
    # PDF Export Methods
    # =========================================================================

    def _export_conversation_pdf(
        self,
        messages: List[Dict[str, Any]],
        title: str,
        character_name: str,
        include_timestamps: bool
    ) -> io.BytesIO:
        """Export conversation to PDF format."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        styles = getSampleStyleSheet()
        story = []

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER
        )

        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            spaceAfter=20,
            alignment=TA_CENTER
        )

        user_style = ParagraphStyle(
            'User',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=10,
            leftIndent=0,
            textColor=colors.HexColor('#1a365d')
        )

        assistant_style = ParagraphStyle(
            'Assistant',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=10,
            leftIndent=20,
            textColor=colors.HexColor('#2d3748')
        )

        timestamp_style = ParagraphStyle(
            'Timestamp',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            spaceBefore=2,
            spaceAfter=8
        )

        # Title
        story.append(Paragraph(title, title_style))

        # Metadata
        meta_text = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} | {len(messages)} messages | Character: {character_name}"
        story.append(Paragraph(meta_text, meta_style))
        story.append(Spacer(1, 20))

        # Messages
        for msg in messages:
            role = msg.get('role', 'user')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')

            # Escape special characters for PDF
            content = self._escape_for_pdf(content)

            if role == 'user':
                story.append(Paragraph(f"<b>You:</b>", user_style))
                story.append(Paragraph(content, user_style))
            else:
                story.append(Paragraph(f"<b>{character_name}:</b>", assistant_style))
                story.append(Paragraph(content, assistant_style))

            if include_timestamps and timestamp:
                story.append(Paragraph(timestamp, timestamp_style))

            story.append(Spacer(1, 10))

        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph("Exported from MinouChat", footer_style))

        doc.build(story)
        buffer.seek(0)
        return buffer

    def _export_summary_pdf(
        self,
        summary_text: str,
        title: str,
        character_name: str,
        key_points: Optional[List[str]],
        original_message_count: int
    ) -> io.BytesIO:
        """Export summary to PDF format."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        styles = getSampleStyleSheet()
        story = []

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER
        )

        # Title
        story.append(Paragraph(title, title_style))

        # Metadata
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            spaceAfter=20,
            alignment=TA_CENTER
        )
        meta_text = f"Summary of {original_message_count} messages | Generated by {character_name}"
        story.append(Paragraph(meta_text, meta_style))
        story.append(Spacer(1, 20))

        # Summary content
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY
        )
        summary_text = self._escape_for_pdf(summary_text)
        story.append(Paragraph(summary_text, body_style))

        # Key points
        if key_points:
            story.append(Spacer(1, 20))
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12
            )
            story.append(Paragraph("Key Points", heading_style))

            bullet_style = ParagraphStyle(
                'Bullet',
                parent=styles['Normal'],
                fontSize=11,
                leftIndent=20,
                spaceAfter=6,
                bulletIndent=10
            )
            for point in key_points:
                point = self._escape_for_pdf(point)
                story.append(Paragraph(f"• {point}", bullet_style))

        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph("Exported from MinouChat", footer_style))

        doc.build(story)
        buffer.seek(0)
        return buffer

    def _escape_for_pdf(self, text: str) -> str:
        """Escape special characters for ReportLab PDF generation."""
        if not text:
            return ""
        # Replace special XML/HTML characters
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        # Replace newlines with break tags
        text = text.replace('\n', '<br/>')
        return text

    # =========================================================================
    # DOCX Export Methods
    # =========================================================================

    def _export_conversation_docx(
        self,
        messages: List[Dict[str, Any]],
        title: str,
        character_name: str,
        include_timestamps: bool
    ) -> io.BytesIO:
        """Export conversation to Word document format."""
        doc = Document()

        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
            f"{len(messages)} messages | Character: {character_name}"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacer

        # Add messages
        for msg in messages:
            role = msg.get('role', 'user')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')

            # Add role header
            role_para = doc.add_paragraph()
            if role == 'user':
                role_run = role_para.add_run("You: ")
                role_run.bold = True
                role_run.font.color.rgb = RGBColor(26, 54, 93)  # Dark blue
            else:
                role_run = role_para.add_run(f"{character_name}: ")
                role_run.bold = True
                role_run.font.color.rgb = RGBColor(45, 55, 72)  # Dark gray

            # Add content
            content_run = role_para.add_run(content)
            content_run.font.size = Pt(11)

            # Add timestamp
            if include_timestamps and timestamp:
                ts_para = doc.add_paragraph()
                ts_run = ts_para.add_run(timestamp)
                ts_run.font.size = Pt(8)
                ts_run.font.color.rgb = RGBColor(128, 128, 128)

            doc.add_paragraph()  # Spacer between messages

        # Add footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Exported from MinouChat")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _export_summary_docx(
        self,
        summary_text: str,
        title: str,
        character_name: str,
        key_points: Optional[List[str]],
        original_message_count: int
    ) -> io.BytesIO:
        """Export summary to Word document format."""
        doc = Document()

        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"Summary of {original_message_count} messages | Generated by {character_name}"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacer

        # Add summary
        doc.add_heading("Summary", level=1)
        summary_para = doc.add_paragraph(summary_text)
        summary_para.paragraph_format.space_after = Pt(12)

        # Add key points
        if key_points:
            doc.add_heading("Key Points", level=1)
            for point in key_points:
                bullet_para = doc.add_paragraph(point, style='List Bullet')

        # Add footer
        doc.add_paragraph()  # Spacer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Exported from MinouChat")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # =========================================================================
    # Markdown Export Methods
    # =========================================================================

    def _export_conversation_markdown(
        self,
        messages: List[Dict[str, Any]],
        title: str,
        character_name: str,
        include_timestamps: bool
    ) -> io.BytesIO:
        """Export conversation to Markdown format."""
        lines = []

        # Title
        lines.append(f"# {title}")
        lines.append("")

        # Metadata
        lines.append(f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
                     f"{len(messages)} messages | Character: {character_name}*")
        lines.append("")
        lines.append("---")
        lines.append("")

        # Messages
        for msg in messages:
            role = msg.get('role', 'user')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')

            if role == 'user':
                lines.append("### You")
            else:
                lines.append(f"### {character_name}")

            lines.append("")
            lines.append(content)
            lines.append("")

            if include_timestamps and timestamp:
                lines.append(f"*{timestamp}*")
                lines.append("")

            lines.append("---")
            lines.append("")

        # Footer
        lines.append("")
        lines.append("*Exported from MinouChat*")

        content = "\n".join(lines)
        buffer = io.BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return buffer

    def _export_summary_markdown(
        self,
        summary_text: str,
        title: str,
        character_name: str,
        key_points: Optional[List[str]],
        original_message_count: int
    ) -> io.BytesIO:
        """Export summary to Markdown format."""
        lines = []

        # Title
        lines.append(f"# {title}")
        lines.append("")

        # Metadata
        lines.append(f"*Summary of {original_message_count} messages | Generated by {character_name}*")
        lines.append("")
        lines.append("---")
        lines.append("")

        # Summary
        lines.append("## Summary")
        lines.append("")
        lines.append(summary_text)
        lines.append("")

        # Key points
        if key_points:
            lines.append("## Key Points")
            lines.append("")
            for point in key_points:
                lines.append(f"- {point}")
            lines.append("")

        # Footer
        lines.append("---")
        lines.append("")
        lines.append("*Exported from MinouChat*")

        content = "\n".join(lines)
        buffer = io.BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return buffer

    # =========================================================================
    # Plain Text Export Methods
    # =========================================================================

    def _export_conversation_text(
        self,
        messages: List[Dict[str, Any]],
        title: str,
        character_name: str,
        include_timestamps: bool
    ) -> io.BytesIO:
        """Export conversation to plain text format."""
        lines = []

        # Title
        lines.append(title)
        lines.append("=" * len(title))
        lines.append("")

        # Metadata
        lines.append(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append(f"Messages: {len(messages)}")
        lines.append(f"Character: {character_name}")
        lines.append("")
        lines.append("-" * 50)
        lines.append("")

        # Messages
        for msg in messages:
            role = msg.get('role', 'user')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')

            if role == 'user':
                lines.append("YOU:")
            else:
                lines.append(f"{character_name.upper()}:")

            lines.append(content)
            lines.append("")

            if include_timestamps and timestamp:
                lines.append(f"[{timestamp}]")
                lines.append("")

            lines.append("-" * 50)
            lines.append("")

        # Footer
        lines.append("")
        lines.append("Exported from MinouChat")

        content = "\n".join(lines)
        buffer = io.BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return buffer

    def _export_summary_text(
        self,
        summary_text: str,
        title: str,
        character_name: str,
        key_points: Optional[List[str]],
        original_message_count: int
    ) -> io.BytesIO:
        """Export summary to plain text format."""
        lines = []

        # Title
        lines.append(title)
        lines.append("=" * len(title))
        lines.append("")

        # Metadata
        lines.append(f"Summary of {original_message_count} messages")
        lines.append(f"Generated by {character_name}")
        lines.append("")
        lines.append("-" * 50)
        lines.append("")

        # Summary
        lines.append("SUMMARY")
        lines.append("-" * 7)
        lines.append("")
        lines.append(summary_text)
        lines.append("")

        # Key points
        if key_points:
            lines.append("KEY POINTS")
            lines.append("-" * 10)
            lines.append("")
            for point in key_points:
                lines.append(f"* {point}")
            lines.append("")

        # Footer
        lines.append("-" * 50)
        lines.append("")
        lines.append("Exported from MinouChat")

        content = "\n".join(lines)
        buffer = io.BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return buffer

    def get_content_type(self, format: ExportFormat) -> str:
        """Get the MIME content type for a format."""
        content_types = {
            ExportFormat.PDF: "application/pdf",
            ExportFormat.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ExportFormat.MARKDOWN: "text/markdown",
            ExportFormat.TEXT: "text/plain"
        }
        return content_types.get(format, "application/octet-stream")

    def get_file_extension(self, format: ExportFormat) -> str:
        """Get the file extension for a format."""
        return format.value

    def generate_document_with_llm(
        self,
        messages: List[Dict[str, Any]],
        document_type: DocumentType,
        export_format: ExportFormat,
        character_name: str,
        character_category: str,
        custom_instructions: Optional[str] = None,
        sfw_mode: bool = True,
        existing_story: Optional[str] = None,
        llm_client: Optional[Any] = None,
        model_config: Optional[Dict[str, Any]] = None
    ) -> Tuple[io.BytesIO, str]:
        """Generate a document using LLM and export to the specified format.

        Args:
            messages: Conversation messages
            document_type: Type of document to generate
            export_format: Output format (PDF, DOCX, MD, TXT)
            character_name: Name of the AI character
            character_category: Category of the persona (Assistant, Coach, etc.)
            custom_instructions: Optional user instructions for the document
            sfw_mode: For Creative/Roleplay, whether to keep content SFW
            existing_story: For Collaborative Story, existing story content to append to
            llm_client: LLM client for generating content
            model_config: Model configuration for the LLM

        Returns:
            Tuple of (BytesIO buffer with file, suggested filename)
        """
        # Validate document type is valid for category
        valid_types = CATEGORY_DOCUMENTS.get(character_category, [])
        if document_type not in valid_types:
            # Fall back to checking all document types
            if document_type not in DOCUMENT_TYPES:
                raise ValueError(f"Invalid document type: {document_type}")

        # Get the prompt template
        prompt_template = DOCUMENT_PROMPTS.get(document_type)
        if not prompt_template:
            raise ValueError(f"No prompt template for document type: {document_type}")

        # Format conversation for the prompt
        conversation_text = self._format_conversation_for_prompt(messages)

        # Build custom instructions text
        custom_text = ""
        if custom_instructions:
            custom_text = f"USER INSTRUCTIONS: {custom_instructions}"

        # Build SFW instruction for creative/roleplay
        sfw_instruction = ""
        if document_type in [DocumentType.STORY_SUMMARY_BRIEF, DocumentType.STORY_SUMMARY_DETAILED]:
            if sfw_mode:
                sfw_instruction = "CONTENT GUIDELINE: Keep the summary appropriate for all audiences (SFW). Focus on plot, character development, and emotional beats rather than explicit content."
            else:
                sfw_instruction = "CONTENT GUIDELINE: Include all relevant story elements including mature themes if present in the original conversation."

        # Fill in the prompt template
        prompt = prompt_template.format(
            conversation=conversation_text,
            character_name=character_name,
            custom_instructions=custom_text,
            sfw_instruction=sfw_instruction
        )

        # Generate content with LLM
        if llm_client is None:
            from .llm_client import llm_client as default_llm_client
            llm_client = default_llm_client

        if model_config is None:
            model_config = {
                "provider": "ollama",
                "model": "llama3.1:8b",
                "temperature": 0.7,
                "max_tokens": 4096
            }

        try:
            llm_messages = [{"role": "user", "content": prompt}]
            generated_content = llm_client.generate_response_with_config(
                messages=llm_messages,
                system_prompt="You are a professional document generator. Create well-structured, clear documents based on the provided template and conversation context.",
                model_config=model_config
            )
        except Exception as e:
            logger.error(f"LLM generation failed: {e}")
            raise ValueError(f"Failed to generate document: {str(e)}")

        # For collaborative story, append to existing content
        if document_type == DocumentType.COLLABORATIVE_STORY and existing_story:
            generated_content = existing_story + "\n\n---\n\n" + generated_content

        # Get document type info for title
        doc_info = DOCUMENT_TYPES.get(document_type)
        doc_name = doc_info.name if doc_info else "Document"
        title = f"{doc_name} - {character_name}"

        # Export to the specified format
        buffer = self._export_generated_content(
            content=generated_content,
            title=title,
            export_format=export_format,
            character_name=character_name,
            document_type=document_type
        )

        # Generate filename
        safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()[:50]
        safe_title = safe_title.replace(' ', '_')
        filename = f"{safe_title}.{self.get_file_extension(export_format)}"

        return buffer, filename

    def _format_conversation_for_prompt(
        self,
        messages: List[Dict[str, Any]],
        max_messages: int = 50
    ) -> str:
        """Format conversation messages for inclusion in an LLM prompt.

        Args:
            messages: List of message dictionaries
            max_messages: Maximum messages to include (most recent)

        Returns:
            Formatted conversation text
        """
        # Take most recent messages if too many
        if len(messages) > max_messages:
            messages = messages[-max_messages:]

        lines = []
        for msg in messages:
            role = msg.get('role', 'user')
            content = msg.get('content', '')

            # Skip system messages
            if role == 'system':
                continue

            # Truncate very long messages
            if len(content) > 2000:
                content = content[:2000] + "... [truncated]"

            role_label = "User" if role == 'user' else "Assistant"
            lines.append(f"{role_label}: {content}")

        return "\n\n".join(lines)

    def _export_generated_content(
        self,
        content: str,
        title: str,
        export_format: ExportFormat,
        character_name: str,
        document_type: DocumentType
    ) -> io.BytesIO:
        """Export LLM-generated content to the specified format.

        Args:
            content: The generated markdown content
            title: Document title
            export_format: Output format
            character_name: Character name for metadata
            document_type: Type of document (for styling)

        Returns:
            BytesIO buffer containing the file
        """
        doc_info = DOCUMENT_TYPES.get(document_type)
        doc_type_name = doc_info.name if doc_info else "Document"

        if export_format == ExportFormat.MARKDOWN:
            return self._export_generated_markdown(content, title, character_name, doc_type_name)
        elif export_format == ExportFormat.TEXT:
            return self._export_generated_text(content, title, character_name, doc_type_name)
        elif export_format == ExportFormat.PDF:
            return self._export_generated_pdf(content, title, character_name, doc_type_name)
        elif export_format == ExportFormat.DOCX:
            return self._export_generated_docx(content, title, character_name, doc_type_name)
        else:
            raise ValueError(f"Unsupported export format: {export_format}")

    def _export_generated_markdown(
        self,
        content: str,
        title: str,
        character_name: str,
        doc_type_name: str
    ) -> io.BytesIO:
        """Export generated content as Markdown."""
        lines = [
            f"# {title}",
            "",
            f"*{doc_type_name} | Generated by {character_name} | {datetime.now().strftime('%Y-%m-%d %H:%M')}*",
            "",
            "---",
            "",
            content,
            "",
            "---",
            "",
            "*Generated by MinouChat*"
        ]
        text = "\n".join(lines)
        buffer = io.BytesIO(text.encode('utf-8'))
        buffer.seek(0)
        return buffer

    def _export_generated_text(
        self,
        content: str,
        title: str,
        character_name: str,
        doc_type_name: str
    ) -> io.BytesIO:
        """Export generated content as plain text."""
        # Remove markdown formatting for plain text
        import re
        plain_content = content
        # Remove headers (## Header -> HEADER)
        plain_content = re.sub(r'^#{1,6}\s+(.+)$', lambda m: m.group(1).upper(), plain_content, flags=re.MULTILINE)
        # Remove bold/italic
        plain_content = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', plain_content)
        # Remove code blocks
        plain_content = re.sub(r'```[\s\S]*?```', '', plain_content)

        lines = [
            title,
            "=" * len(title),
            "",
            f"{doc_type_name}",
            f"Generated by {character_name}",
            f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "",
            "-" * 50,
            "",
            plain_content,
            "",
            "-" * 50,
            "",
            "Generated by MinouChat"
        ]
        text = "\n".join(lines)
        buffer = io.BytesIO(text.encode('utf-8'))
        buffer.seek(0)
        return buffer

    def _export_generated_pdf(
        self,
        content: str,
        title: str,
        character_name: str,
        doc_type_name: str
    ) -> io.BytesIO:
        """Export generated content as PDF."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        styles = getSampleStyleSheet()
        story = []

        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER
        )

        # Meta style
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            spaceAfter=20,
            alignment=TA_CENTER
        )

        # Body style
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_LEFT
        )

        # Heading styles
        h2_style = ParagraphStyle(
            'H2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=16,
            spaceAfter=8
        )

        # Title
        story.append(Paragraph(self._escape_for_pdf(title), title_style))

        # Metadata
        meta_text = f"{doc_type_name} | Generated by {character_name} | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        story.append(Paragraph(meta_text, meta_style))
        story.append(Spacer(1, 20))

        # Parse and render markdown content
        import re
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue

            # Check for headers
            if line.startswith('## '):
                header_text = self._escape_for_pdf(line[3:])
                story.append(Paragraph(header_text, h2_style))
            elif line.startswith('# '):
                # Skip top-level headers (we already have title)
                pass
            elif line.startswith('- ') or line.startswith('* '):
                bullet_text = self._escape_for_pdf(line[2:])
                bullet_style = ParagraphStyle(
                    'Bullet',
                    parent=body_style,
                    leftIndent=20,
                    bulletIndent=10
                )
                story.append(Paragraph(f"• {bullet_text}", bullet_style))
            else:
                # Regular paragraph - remove markdown formatting
                clean_line = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', line)
                story.append(Paragraph(self._escape_for_pdf(clean_line), body_style))

        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph("Generated by MinouChat", footer_style))

        doc.build(story)
        buffer.seek(0)
        return buffer

    def _export_generated_docx(
        self,
        content: str,
        title: str,
        character_name: str,
        doc_type_name: str
    ) -> io.BytesIO:
        """Export generated content as Word document."""
        doc = Document()

        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"{doc_type_name} | Generated by {character_name} | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacer

        # Parse and render markdown content
        import re
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check for headers
            if line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                # Skip top-level headers
                pass
            elif line.startswith('- ') or line.startswith('* '):
                bullet_text = line[2:]
                # Remove markdown formatting
                bullet_text = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', bullet_text)
                doc.add_paragraph(bullet_text, style='List Bullet')
            else:
                # Regular paragraph - remove markdown formatting
                clean_line = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', line)
                doc.add_paragraph(clean_line)

        # Footer
        doc.add_paragraph()  # Spacer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Generated by MinouChat")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


# Global service instance
export_service = ExportService()
